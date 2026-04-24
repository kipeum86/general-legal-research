"""
Consumer adapter for citation-auditor output in DOCX deliverables.

Responsibilities:
  1. `inject_unverified_tags(body_markdown, aggregated)` — returns the markdown
     string with `[Unverified]` / `[Partially Unverified]` tags inserted at the
     end of each failing claim's sentence. Processes in reverse offset order so
     earlier positions remain stable.
  2. `append_citation_audit_log(doc, aggregated)` — appends a
     `부록: 검증 로그 (Citation Audit Log)` heading + audit table to an existing
     `python-docx` Document, using this project's CJK-safe font conventions.
  3. `load_aggregated(path)` — convenience loader for the JSON produced by
     `python3 -m citation_auditor aggregate`.

Kept deliberately separate from `citation_auditor/` so the vendored package is
never modified. See CLAUDE.md §5 Step 7 / Step 9 for workflow integration.
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


FONT_BODY = "Times New Roman"
FONT_BODY_KO = "맑은 고딕"
FONT_SIZE = Pt(10.5)
FONT_SIZE_SMALL = Pt(9)
FONT_SIZE_HEADING = Pt(13)

COLOR_NAVY = RGBColor(0x19, 0x37, 0x6D)
COLOR_ALERT = RGBColor(0xAA, 0x14, 0x14)
COLOR_BODY = RGBColor(0x00, 0x00, 0x00)
COLOR_MUTED = RGBColor(0x60, 0x60, 0x60)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FILL_HEADER = "1B2A4A"
FILL_UNVERIFIED = "FFF8E1"


TAG_CONTRADICTED = " [Unverified]"
TAG_UNKNOWN = " [Partially Unverified]"


def _set_run_font(run, *, size=None, bold=False, color=None):
    """Apply CJK-safe font (Times New Roman + 맑은 고딕 eastAsia) to a run.

    Mirrors `scripts/render_acp_comparison_docx.py::_set_run_font` so audit
    output visually matches the rest of the document.
    """
    run.font.name = FONT_BODY
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:ascii"), FONT_BODY)
    rf.set(qn("w:hAnsi"), FONT_BODY)
    rf.set(qn("w:eastAsia"), FONT_BODY_KO)
    run.font.size = size or FONT_SIZE
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _clear_cell(cell) -> None:
    cell.text = ""


def _verdict_label_tag(label: str) -> str:
    if label == "contradicted":
        return TAG_CONTRADICTED
    if label in {"unknown", "unsupported", "source_unavailable", "verifier_unavailable"}:
        return TAG_UNKNOWN
    return ""


def load_aggregated(path: str | Path) -> dict[str, Any]:
    """Read aggregated.json produced by `python3 -m citation_auditor aggregate`."""
    return json.loads(Path(path).read_text(encoding="utf-8"))


def inject_unverified_tags_with_report(body_markdown: str, aggregated: dict[str, Any]) -> tuple[str, dict[str, Any]]:
    """Insert failing-claim tags and report span validation outcomes.

    The `aggregated` dict is the schema emitted by `citation_auditor`'s
    AggregateOutput model: {"aggregated": [{"claim": {...}, "verdict": {...}}, ...]}.
    `sentence_span.end` is an absolute character offset into the original body.

    Insertions are processed in descending `end` order so that earlier offsets
    remain valid after prior insertions.
    """
    items = list(aggregated.get("aggregated", []))
    skip_ranges = _skip_ranges(body_markdown)
    insertions: list[tuple[int, str]] = []
    skipped: list[dict[str, Any]] = []
    relocated = 0
    offset_only = 0

    for item in items:
        label = item.get("verdict", {}).get("label")
        tag = _verdict_label_tag(label)
        if not tag:
            continue
        claim = item.get("claim", {}) or {}
        claim_text = str(claim.get("text") or "")
        span = claim.get("sentence_span", {}) or {}
        start = span.get("start")
        end = span.get("end")
        if not isinstance(start, int) or not isinstance(end, int):
            skipped.append(_skip_record(claim_text, "missing_or_invalid_span"))
            continue
        if start < 0 or end < start or end > len(body_markdown):
            skipped.append(_skip_record(claim_text, "span_out_of_bounds", start=start, end=end))
            continue
        if _span_overlaps_skip_range(start, end, skip_ranges):
            skipped.append(_skip_record(claim_text, "span_inside_code_or_quote", start=start, end=end))
            continue

        insertion_end = end
        if claim_text:
            span_text = body_markdown[start:end]
            if span_text != claim_text:
                matches = [match.start() for match in re.finditer(re.escape(claim_text), body_markdown)]
                if len(matches) == 1:
                    insertion_end = matches[0] + len(claim_text)
                    relocated += 1
                elif len(matches) > 1:
                    skipped.append(_skip_record(claim_text, "ambiguous_claim_text", start=start, end=end))
                    continue
                else:
                    skipped.append(_skip_record(claim_text, "claim_text_not_found", start=start, end=end))
                    continue
        else:
            offset_only += 1
        insertions.append((insertion_end, tag))

    insertions.sort(key=lambda pair: pair[0], reverse=True)
    out = body_markdown
    for end, tag in insertions:
        if 0 <= end <= len(out):
            out = out[:end] + tag + out[end:]
    report = {
        "failing_claims": len(insertions) + len(skipped),
        "inserted": len(insertions),
        "relocated": relocated,
        "offset_only": offset_only,
        "skipped": skipped,
    }
    return out, report


def inject_unverified_tags(body_markdown: str, aggregated: dict[str, Any]) -> str:
    """Backward-compatible wrapper returning only the annotated markdown."""
    annotated, _report = inject_unverified_tags_with_report(body_markdown, aggregated)
    return annotated


def _skip_record(claim_text: str, reason: str, **extra: Any) -> dict[str, Any]:
    record = {"claim": claim_text, "reason": reason}
    record.update(extra)
    return record


def _skip_ranges(md_text: str) -> list[tuple[int, int]]:
    ranges: list[tuple[int, int]] = []
    fence_start: int | None = None
    fence_token: str | None = None
    cursor = 0
    for line in md_text.splitlines(keepends=True):
        stripped = line.lstrip()
        if fence_start is None:
            match = re.match(r"([`~]{3,})", stripped)
            if match:
                fence_start = cursor
                fence_token = match.group(1)[0] * len(match.group(1))
        else:
            if fence_token is not None and stripped.startswith(fence_token):
                ranges.append((fence_start, cursor + len(line)))
                fence_start = None
                fence_token = None
        cursor += len(line)
    if fence_start is not None:
        ranges.append((fence_start, len(md_text)))

    in_quote = False
    quote_start = 0
    cursor = 0
    for line in md_text.splitlines(keepends=True):
        stripped = line.lstrip()
        if stripped.startswith(">"):
            if not in_quote:
                quote_start = cursor
                in_quote = True
        elif in_quote:
            ranges.append((quote_start, cursor))
            in_quote = False
        cursor += len(line)
    if in_quote:
        ranges.append((quote_start, len(md_text)))
    return ranges


def _span_overlaps_skip_range(start: int, end: int, skip_ranges: list[tuple[int, int]]) -> bool:
    for skip_start, skip_end in skip_ranges:
        if start < skip_end and end > skip_start:
            return True
    return False


def _label_display(label: str) -> str:
    return {
        "verified": "✓ verified",
        "contradicted": "⚠ contradicted",
        "unknown": "? unknown",
    }.get(label, label or "—")


def _evidence_string(evidence: list[dict[str, Any]]) -> str:
    urls = [e.get("url", "").strip() for e in evidence if e.get("url")]
    return "; ".join(u for u in urls if u) or "—"


def append_citation_audit_log(
    doc: Document,
    aggregated: dict[str, Any],
    *,
    audit_status: str = "complete",
    artifact_path: str | Path | None = None,
    resolution_message: str | None = None,
    insertion_report: dict[str, Any] | None = None,
) -> None:
    """Append the audit appendix (heading + table + disclaimer) to a Document.

    The heading is a plain paragraph (no Word built-in Heading style is used, to
    avoid inheriting style fonts that override the CJK eastAsia setting).
    """
    items = list(aggregated.get("aggregated", []))

    doc.add_paragraph()

    heading = doc.add_paragraph()
    hr = heading.add_run("부록: 검증 로그 (Citation Audit Log)")
    _set_run_font(hr, size=FONT_SIZE_HEADING, bold=True, color=COLOR_NAVY)

    _append_status_block(
        doc,
        audit_status=audit_status,
        artifact_path=artifact_path,
        resolution_message=resolution_message,
        insertion_report=insertion_report,
    )

    if not items:
        note = doc.add_paragraph()
        nr = note.add_run("감사 대상 클레임이 없습니다." if audit_status == "complete" else "인용 감사 결과가 최종 문서에 포함되지 않았습니다.")
        _set_run_font(nr, size=FONT_SIZE_SMALL, color=COLOR_MUTED)
        return

    table = doc.add_table(rows=1 + len(items), cols=5)
    table.autofit = True

    headers = ["#", "클레임 (Claim)", "판정 (Verdict)", "Verifier", "근거 (Evidence)"]
    header_row = table.rows[0].cells
    for j, header in enumerate(headers):
        cell = header_row[j]
        _clear_cell(cell)
        _set_cell_shading(cell, FILL_HEADER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = cell.paragraphs[0].add_run(header)
        _set_run_font(run, size=FONT_SIZE_SMALL, bold=True, color=COLOR_WHITE)

    for i, item in enumerate(items, start=1):
        claim_text = item.get("claim", {}).get("text", "")
        verdict = item.get("verdict", {}) or {}
        label = verdict.get("label", "")
        verifier = verdict.get("verifier_name", "") or ""
        evidence_str = _evidence_string(verdict.get("evidence", []) or [])

        row = table.rows[i].cells
        _clear_cell(row[0])
        run = row[0].paragraphs[0].add_run(str(i))
        _set_run_font(run, size=FONT_SIZE_SMALL)

        _clear_cell(row[1])
        run = row[1].paragraphs[0].add_run(claim_text)
        _set_run_font(run, size=FONT_SIZE_SMALL)

        _clear_cell(row[2])
        if label == "contradicted":
            _set_cell_shading(row[2], FILL_UNVERIFIED)
        run = row[2].paragraphs[0].add_run(_label_display(label))
        _set_run_font(
            run,
            size=FONT_SIZE_SMALL,
            bold=(label != "verified"),
            color=(COLOR_ALERT if label == "contradicted" else COLOR_BODY),
        )

        _clear_cell(row[3])
        run = row[3].paragraphs[0].add_run(verifier)
        _set_run_font(run, size=FONT_SIZE_SMALL, color=COLOR_MUTED)

        _clear_cell(row[4])
        run = row[4].paragraphs[0].add_run(evidence_str)
        _set_run_font(run, size=FONT_SIZE_SMALL, color=COLOR_MUTED)

    disc = doc.add_paragraph()
    dr = disc.add_run("※ 자동 감사 결과는 참고 자료이며 전문 변호사의 검토를 대체하지 않습니다.")
    _set_run_font(dr, size=FONT_SIZE_SMALL, color=COLOR_MUTED)


def _append_status_block(
    doc: Document,
    *,
    audit_status: str,
    artifact_path: str | Path | None,
    resolution_message: str | None,
    insertion_report: dict[str, Any] | None,
) -> None:
    lines = [f"Audit status: {audit_status}"]
    if artifact_path is not None:
        lines.append(f"Audit artifact: {artifact_path}")
    if resolution_message:
        lines.append(f"Artifact resolution: {resolution_message}")
    if insertion_report:
        inserted = insertion_report.get("inserted", 0)
        skipped = len(insertion_report.get("skipped", []) or [])
        relocated = insertion_report.get("relocated", 0)
        lines.append(f"Inline tag insertion: inserted={inserted}, relocated={relocated}, skipped={skipped}")

    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        _set_run_font(run, size=FONT_SIZE_SMALL, color=COLOR_MUTED)
