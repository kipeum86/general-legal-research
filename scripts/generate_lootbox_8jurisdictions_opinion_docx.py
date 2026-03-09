from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT_KO = "Malgun Gothic"
COLOR_PRIMARY = RGBColor(25, 55, 109)
COLOR_SECONDARY = RGBColor(68, 84, 106)
COLOR_MUTED = RGBColor(89, 89, 89)
COLOR_ALERT = RGBColor(192, 0, 0)
AS_OF_DATE = date(2026, 3, 5)


def set_font(run, name=FONT_KO, size_pt=None, bold=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill="DCE6F1"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, instruction):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    paragraph._p.append(fld)


def style_normal(document):
    normal = document.styles["Normal"]
    normal.font.name = FONT_KO
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = document.styles[style_name]
        style.font.name = FONT_KO
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = COLOR_PRIMARY
        style.paragraph_format.space_before = Pt(10 if style_name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    list_bullet = document.styles["List Bullet"]
    list_bullet.font.name = FONT_KO
    list_bullet._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    list_bullet.font.size = Pt(10.5)

    list_number = document.styles["List Number"]
    list_number.font.name = FONT_KO
    list_number._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    list_number.font.size = Pt(10.0)


def setup_header_footer(section):
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("JINJU LAW FIRM | Game Legal Regulatory Research")
    set_font(hr, size_pt=9, color=RGBColor(120, 120, 120))

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Page ")
    set_font(fr, size_pt=9, color=RGBColor(120, 120, 120))
    add_field(fp, "PAGE")
    fr2 = fp.add_run(" of ")
    set_font(fr2, size_pt=9, color=RGBColor(120, 120, 120))
    add_field(fp, "NUMPAGES")


def add_heading(document, text, level=1):
    p = document.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size_pt=16 if level == 1 else (13 if level == 2 else 11), bold=True, color=COLOR_PRIMARY)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(document, text, bold=False, size=10.5, color=None):
    p = document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size_pt=size, bold=bold, color=color)
    return p


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size_pt=10.5)
    return p


def add_cover(document):
    today = date.today()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("비공개 / 클라이언트 내부검토용")
    set_font(run, size_pt=11, bold=True, color=COLOR_ALERT)

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("로트박스(Loot Box) 규제 8개국 비교 의견서")
    set_font(run, size_pt=24, bold=True, color=COLOR_PRIMARY)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("벨기에 · 중국 · 세르비아 · 슬로바키아 · 네덜란드 · 대한민국 · 대만 · 일본")
    set_font(run, size_pt=12, bold=False, color=COLOR_SECONDARY)

    document.add_paragraph()
    info_table = document.add_table(rows=3, cols=2)
    info_table.style = "Table Grid"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    info_rows = [
        ("의뢰인", "한국 게임사 클라이언트"),
        ("기준일", f"{AS_OF_DATE.year}년 {AS_OF_DATE.month}월 {AS_OF_DATE.day}일"),
        ("작성일", f"{today.year}년 {today.month}월 {today.day}일"),
    ]
    for r_idx, (label, value) in enumerate(info_rows):
        label_cell = info_table.rows[r_idx].cells[0]
        value_cell = info_table.rows[r_idx].cells[1]
        label_cell.width = Cm(4.0)
        value_cell.width = Cm(9.5)
        for cell in (label_cell, value_cell):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
        shade_cell(label_cell, "E8EEF7")
        lp = label_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(label)
        set_font(lr, size_pt=10.5, bold=True, color=COLOR_PRIMARY)
        vp = value_cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vr = vp.add_run(value)
        set_font(vr, size_pt=10.5)

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.right_indent = Cm(0.2)
    run = p.add_run(
        "면책: 본 문서는 공개된 법령·행정자료 기반의 리서치 의견서이며, 특정 사건에 대한 확정적 법률의견(법률자문)으로 대체되지 않습니다. "
        "국가별 집행관행과 최신 개정은 수시로 변동될 수 있으므로 실제 출시 전 현지 변호사 검토가 필요합니다."
    )
    set_font(run, size_pt=9.5, color=COLOR_MUTED)

    document.add_page_break()


def add_executive_summary(document):
    add_heading(document, "1. 집행요약", level=1)
    add_body(
        document,
        "질문: 벨기에, 중국, 세르비아, 슬로바키아, 네덜란드, 대한민국, 대만, 일본에서 로트박스 판매가 금지되는지 여부."
    )
    add_body(
        document,
        f"핵심 결론(기준일: {AS_OF_DATE.isoformat()}): 8개국 중 전면적/실질적 금지에 가장 가까운 국가는 벨기에이며, "
        "일본은 컴프가챠(complete gacha) 유형에 한정된 부분 금지, 나머지 국가는 주로 확률공개·소비자보호·도박법 일반요건에 따른 조건부 규제 체계입니다."
    )
    add_bullet(document, "벨기에: 유료 로트박스는 도박법상 게임요건 충족 시 사실상 판매 금지(고위험).")
    add_bullet(document, "일본: 일반 유료 가챠는 전면 금지 아님. 다만 컴프가챠는 위법 리스크가 매우 높음.")
    add_bullet(document, "중국·대한민국·대만: 금지보다는 확률공개 및 운영투명성 의무 중심.")
    add_bullet(document, "네덜란드·세르비아·슬로바키아: 전면 금지 아님. 경제적 가치/현금화 가능성/독립 도박성 여부를 중심으로 사안별 판단.")


def add_conclusion_table(document):
    add_heading(document, "2. 국가별 결론 표", level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    repeat_table_header(table.rows[0])
    col_widths = [2.2, 4.1, 2.6, 6.3]

    headers = ["국가", "결론(금지 여부)", "실무 리스크 등급", "출시 권고"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Cm(col_widths[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        cell.text = text
        shade_cell(cell, "E8EEF7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.runs[0], size_pt=10.5, bold=True, color=COLOR_PRIMARY)

    rows = [
        ("벨기에", "사실상 금지(유료 로트박스)", "매우 높음", "벨기에 빌드에서 유료 랜덤 판매 비활성화"),
        ("중국", "전면 금지 아님", "중간~높음", "확률·결과 공개, 로그보관, 미성년자 규제 준수"),
        ("세르비아", "전면 금지 아님", "중간", "추가 결제 없는 랜덤보상 중심 설계, 현금화 차단"),
        ("슬로바키아", "전면 금지 아님", "중간", "스테이크/환금성 구조 회피, 사전 법률검토"),
        ("네덜란드", "전면 금지 아님(사안별)", "중간~높음", "거래가능 아이템·독립 도박성 구조 회피"),
        ("대한민국", "전면 금지 아님", "중간~높음", "확률형 아이템 표시·검증 체계 상시 운영"),
        ("대만", "전면 금지 아님", "중간", "확률값 및 변경 7일 전 고지 체계 구축"),
        ("일본", "부분 금지(컴프가챠)", "높음", "컴프가챠 구조 금지, 일반 가챠는 공시·자율규제 준수"),
    ]

    for row_data in rows:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            row[i].width = Cm(col_widths[i])
            row[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(row[i])
            row[i].text = text
            p = row[i].paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 3 else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.runs[0], size_pt=10)


def add_country_sections(document):
    add_heading(document, "3. 국가별 상세 분석", level=1)

    sections = [
        {
            "title": "3.1 벨기에(Belgium)",
            "status": "결론: 유료 로트박스는 사실상 금지(매우 높은 집행 리스크).",
            "analysis": [
                "벨기에 게임위원회(Belgian Gaming Commission) 보고서는 EA/Blizzard/Valve 등 주요 게임의 유료 로트박스가 도박법상 게임요건을 충족한다고 보았습니다.",
                "해당 보고서는 '유료 로트박스를 게임 내에서 금지(prohibited)'하고, 라이선스 없이 운영되는 경우 중단 필요성을 명시합니다.",
                "벨기에 도박법은 대가(consideration), 우연성(chance), 이익/손실(outcome)을 포함하는 구조를 폭넓게 포섭하므로, 유료 랜덤형 BM은 보수적으로 차단하는 것이 안전합니다.",
            ],
            "action": [
                "벨기에 서비스에서는 유료 랜덤뽑기 판매를 원칙적으로 비활성화.",
                "대체 BM(직접구매, 시즌패스, 정액형 보상)을 우선 설계.",
            ],
            "sources": ["[P1]", "[P2]"],
        },
        {
            "title": "3.2 중국(China)",
            "status": "결론: 전면 금지 아님(강한 운영규제).",
            "analysis": [
                "중국 문화부 고시는 온라인게임의 랜덤추첨 제공 시 아이템명·속성·추첨확률을 사전 또는 눈에 띄는 방식으로 공시하고, 결과 공표 및 기록 보관을 요구합니다.",
                "즉, 로트박스 자체를 일률 금지하기보다는 확률정보 투명성·사후검증 가능성을 강하게 요구하는 구조입니다.",
                "별도로 국가신문출판서는 미성년자 게임시간 제한 등 강한 청소년 보호규제를 집행하고 있어 운영정책(로그인, 시간제한, 결제통제)과 함께 설계해야 합니다.",
            ],
            "action": [
                "중국어 UI에서 확률정보 고정 노출 + 변경이력 관리.",
                "랜덤결과 로그 보존체계 및 규제기관 제출 가능 포맷 확보.",
            ],
            "sources": ["[P3]", "[P4]"],
        },
        {
            "title": "3.3 세르비아(Serbia)",
            "status": "결론: 전면 금지 아님(사안별 도박성 판단).",
            "analysis": [
                "세르비아 게임법은 대가성(stake), 우연성, 당첨(winnings)을 중심으로 도박을 정의합니다.",
                "세르비아 조세당국 FAQ는 '게임 판매가격에 포함되어 추가요금 없이 제공되는 미스터리 박스'는 원칙적으로 경품성 게임으로 보기 어렵다는 취지의 해석을 제시합니다.",
                "다만 별도 결제, 환금성, 교환시장 연계가 결합되면 도박법 재분류 리스크가 올라갑니다.",
            ],
            "action": [
                "세르비아 대상 빌드에서 유료 재화로의 직접 랜덤추첨 구매를 최소화.",
                "아이템 현금화/2차 거래 연계를 차단.",
            ],
            "sources": ["[P5]", "[P6]"],
        },
        {
            "title": "3.4 슬로바키아(Slovakia)",
            "status": "결론: 전면 금지 아님(일반 도박법 요건 충족 시 위험).",
            "analysis": [
                "슬로바키아 도박법(Act No. 30/2019)은 플레이어의 베팅, 우연성, 승리/당첨 요소를 결합한 구조를 도박으로 봅니다.",
                "현행 법체계에서 로트박스를 명시적으로 전면 금지하는 특별조항은 확인되지 않았습니다.",
                "따라서 환금성 있는 가상재화, 제3자 거래시장, 별도 베팅구조 여부에 따라 규제당국 해석이 달라질 수 있습니다.",
            ],
            "action": [
                "슬로바키아 출시 시 환금성 약관 금지조항 및 기술적 차단 병행.",
                "중요 업데이트(거래기능 추가 등) 전 현지 자문 재검토.",
            ],
            "sources": ["[P7]"],
        },
        {
            "title": "3.5 네덜란드(Netherlands)",
            "status": "결론: 전면 금지 아님(케이스별 고위험).",
            "analysis": [
                "네덜란드 도박법은 원칙적으로 라이선스 없는 우연성 게임 제공을 금지합니다.",
                "다만 최고행정법원(2022)은 EA FIFA 팩 관련 제재 사건에서 해당 팩을 독립적 게임으로 보기 어렵다는 취지로 감독당국 제재를 취소했습니다.",
                "네덜란드 감독당국(KSA)은 현재도 로트박스가 도박정의에 해당하는지 개별 판단 대상임을 명시하고 있어, 구조 설계가 핵심입니다.",
            ],
            "action": [
                "거래가능(transferable) 아이템·현금화 경로 제거.",
                "패키지 개봉을 독립 게임처럼 보이게 하는 UX/BM(별도 베팅형)을 회피.",
            ],
            "sources": ["[P8]", "[P9]", "[P10]"],
        },
        {
            "title": "3.6 대한민국(South Korea)",
            "status": "결론: 전면 금지 아님(확률공시 중심의 강한 규제).",
            "analysis": [
                "게임산업진흥에 관한 법률은 확률형 아이템의 종류·효과·획득확률 표시의무(제33조의2)를 도입·강화했습니다.",
                "문화체육관광부는 2024년 3월 22일 시행을 공식 발표했고, 확률정보 미표시·허위표시에 대한 제재 프레임이 작동 중입니다.",
                "즉, 한국은 로트박스 금지국이 아니라 '정확한 확률공시 + 사후 입증가능성'을 강제하는 관할입니다.",
            ],
            "action": [
                "확률표시 데이터와 실제 서버확률 간 상시 대사(automated reconciliation).",
                "이벤트/픽업 변경 시 표시값 동기화 배포체계 의무화.",
            ],
            "sources": ["[P11]", "[P12]"],
        },
        {
            "title": "3.7 대만(Taiwan)",
            "status": "결론: 전면 금지 아님(계약규정상 확률공개 의무).",
            "analysis": [
                "대만 '온라인게임 서비스 표준계약 필수·금지 기재사항' 제6점은 확률형 상품에 대해 모든 확률값을 공시하고, 확률 변경 시 7일 전 공지를 요구합니다.",
                "이는 판매 자체 금지보다 정보비대칭 해소(소비자보호) 중심의 규제 방식입니다.",
                "실무상 약관·게임 내 공지·고객센터 문구를 일치시켜 계약분쟁 리스크를 낮추는 것이 중요합니다.",
            ],
            "action": [
                "대만 번체 UI 확률표시 + 변경 7일 전 고지 워크플로우 구축.",
                "고객문의 대응 스크립트(확률 산정 기준, 변경 이력) 표준화.",
            ],
            "sources": ["[P13]"],
        },
        {
            "title": "3.8 일본(Japan)",
            "status": "결론: 전면 금지 아님, 단 컴프가챠는 사실상 금지(부분 금지).",
            "analysis": [
                "일본 소비자청 자료는 '유료 가챠 자체'는 즉시 위법으로 보지 않으면서도, 특정 완성형 컴프가챠 구조는 경품표시법상 문제될 수 있음을 명시합니다.",
                "업계단체(CESA)는 확률표시 등 자율규제 기준을 지속 공표하고 있으며, 대형 사업자는 통상 이를 준수합니다.",
                "따라서 일본은 BM 구조별로 허용/위험이 갈리는 관할로, 컴프가챠형 완성보상 설계는 금지하는 것이 안전합니다.",
            ],
            "action": [
                "세트완성 강제형 가챠(컴프가챠) 설계 금지.",
                "일본어 확률공시 기준을 국내·글로벌 빌드와 분리 관리.",
            ],
            "sources": ["[P14]", "[P15]"],
        },
    ]

    for sec in sections:
        add_heading(document, sec["title"], level=2)
        p = add_body(document, sec["status"], bold=True)
        p.paragraph_format.space_after = Pt(4)
        for line in sec["analysis"]:
            add_bullet(document, line)

        add_body(document, "실무 권고:", bold=True)
        for line in sec["action"]:
            add_bullet(document, line)
        add_body(document, f"주요 근거: {', '.join(sec['sources'])}", size=9.8, color=COLOR_SECONDARY)


def add_compliance_framework(document):
    add_heading(document, "4. 한국 게임사 실행 프레임워크", level=1)
    add_body(
        document,
        "아래 체크리스트는 8개국 동시 서비스 기준의 보수적(규제친화적) 최소선입니다. 벨기에·일본은 별도 게이팅(기능 차단)이 필요합니다."
    )

    add_heading(document, "4.1 글로벌 공통 최소요건", level=2)
    items = [
        "확률형 아이템 테이블(아이템별/등급별/기간별) 버전관리 및 배포 로그 보존",
        "확률표시값과 서버 실제값 일치 검증(정기 샘플링 + 자동 대사)",
        "중복 아이템 보상정책(대체재/파편/천장 시스템) 문서화",
        "약관·인게임 공지·CS 응답문구 간 표현 불일치 금지",
        "거래가능/환금가능 아이템의 지역별 차단 정책(특히 벨기에·네덜란드)",
    ]
    for text in items:
        add_bullet(document, text)

    add_heading(document, "4.2 국가별 게이팅 규칙(권장)", level=2)
    gate_table = document.add_table(rows=1, cols=3)
    gate_table.style = "Table Grid"
    gate_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    gate_table.autofit = False
    repeat_table_header(gate_table.rows[0])
    col_widths = [3.4, 4.6, 7.1]
    heads = ["국가군", "핵심 게이팅", "운영팀 액션"]
    for i, h in enumerate(heads):
        cell = gate_table.rows[0].cells[i]
        cell.width = Cm(col_widths[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        cell.text = h
        shade_cell(cell, "E8EEF7")
        set_font(cell.paragraphs[0].runs[0], size_pt=10.5, bold=True, color=COLOR_PRIMARY)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    data = [
        ("벨기에", "유료 랜덤 판매 OFF", "직접구매형 상품군으로 대체"),
        ("일본", "컴프가챠형 보상 OFF", "일본 전용 BM 검수 체크리스트 적용"),
        ("네덜란드", "거래/환금 경로 최소화", "경제시스템 변경 시 법무 사전검토"),
        ("중국·한국·대만", "확률공시·변경고지 강화", "표시값-실값 증빙체계 상시화"),
        ("세르비아·슬로바키아", "별도 베팅성 구조 회피", "현지 자문 트리거(업데이트 시)"),
    ]
    for row_data in data:
        row = gate_table.add_row().cells
        for i, text in enumerate(row_data):
            row[i].width = Cm(col_widths[i])
            row[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(row[i])
            row[i].text = text
            row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            row[i].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            row[i].paragraphs[0].paragraph_format.line_spacing = 1.15
            set_font(row[i].paragraphs[0].runs[0], size_pt=10)


def add_source_list(document):
    add_heading(document, "5. 출처 목록(Primary 중심)", level=1)
    sources = [
        "[P1] Belgian Gaming Commission, Research Report on Loot Boxes (official PDF): https://www.gamingcommission.be/sites/default/files/2021-08/onderzoeksrapport-loot-boxen-Engels-publicatie.pdf",
        "[P2] Belgium Gaming Act (7 May 1999, consolidated): https://www.ejustice.just.fgov.be/eli/loi/1999/05/07/1999010222/justel",
        "[P3] China Ministry of Culture Notice (Wen Shi Fa [2016] No.32): https://www.gov.cn/gongbao/content/2017/content_5174528.htm",
        "[P4] National Press and Publication Administration Notice (minor protection, 2021): https://www.nppa.gov.cn/xxfb/tzgs/202108/t20210830_6154.html",
        "[P5] Serbia Law on Games of Chance (Official Gazette portal): https://www.pravno-informacioni-sistem.rs/SlGlasnikPortal/eli/rep/sgrs/skupstina/zakon/2020/18/2/reg",
        "[P6] Serbia Tax Administration FAQ (mystery box Q&A): https://www.purs.gov.rs/sr/fizicka-lica/pregled-propisa/najcesca-pitanja/igre-na-srecu-i-nagradne-igre.html",
        "[P7] Slovakia Act No. 30/2019 on Gambling Games (Slov-Lex): https://www.slov-lex.sk/pravne-predpisy/SK/ZZ/2019/30/20240101.html",
        "[P8] Netherlands Betting and Gaming Act (Wet op de kansspelen): https://wetten.overheid.nl/BWBR0002469/2025-07-01",
        "[P9] Netherlands Council of State ruling (ECLI:NL:RVS:2022:690): https://content.overheid.nl/documenten/eb95eeec-b6f7-463f-95f3-e670f08576ad.pdf",
        "[P10] Dutch Gambling Authority (KSA) – Loot boxes guidance: https://kansspelautoriteit.nl/aanpak-misstanden/aanpak-illegale-kansspelen/gaming/loot-boxes/",
        "[P11] Korea Game Industry Promotion Act: https://law.go.kr/법령/게임산업진흥에관한법률",
        "[P12] MCST press release on probability-item disclosure enforcement (2024-03-22): https://www.mcst.go.kr/kor/s_notice/press/pressView.jsp?pSeq=20857",
        "[P13] Taiwan online game standard contract mandatory/prohibited terms (Art.6 probability disclosure): https://law.moj.gov.tw/LawClass/LawAll.aspx?pcode=J0080089",
        "[P14] Japan CAA material on complete gacha (2012): https://www.caa.go.jp/policies/policy/representation/fair_labeling/pdf/120518premiums_1.pdf",
        "[P15] CESA guideline/industry self-regulation materials: https://www.cesa.or.jp/info/efforts.html",
    ]
    for src in sources:
        if "] " in src:
            code, body = src.split("] ", 1)
            code = f"{code}]"
        else:
            code, body = "[Ref]", src
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p.add_run(f"{code} ")
        set_font(r1, size_pt=9.8, bold=True, color=COLOR_PRIMARY)
        r2 = p.add_run(body)
        set_font(r2, size_pt=9.8)

    add_body(
        document,
        "주: 세르비아·슬로바키아 조문은 원문(현지어) 기준으로 해석했으며, 영어 번역은 비공식 번역/자동번역을 병행 참고했습니다.",
        size=9.5,
        color=COLOR_MUTED,
    )


def main():
    out_path = Path("game-legal-research/output/reports/lootbox_8countries_legal_opinion_ko_2026-03-05.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style_normal(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    setup_header_footer(section)

    add_cover(doc)
    add_executive_summary(doc)
    add_conclusion_table(doc)
    add_country_sections(doc)
    add_compliance_framework(doc)
    add_source_list(doc)

    doc.save(out_path)
    print(str(out_path))


if __name__ == "__main__":
    main()
