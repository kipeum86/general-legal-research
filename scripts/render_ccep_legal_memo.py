"""
Children's Content Engagement Platform — Comprehensive Legal Risk Review
Generates a formal legal memorandum in .docx format.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Fonts & colours ──────────────────────────────────────────────
FONT_BODY = "Times New Roman"
FONT_SIZE = 11
FONT_SIZE_SMALL = 9
COLOR_TITLE = RGBColor(25, 55, 109)
COLOR_MUTED = RGBColor(96, 96, 96)
COLOR_ALERT = RGBColor(170, 20, 20)
COLOR_GREEN = RGBColor(0, 100, 0)
COLOR_AMBER = RGBColor(180, 120, 0)
TABLE_HEADER_FILL = "19376D"
TABLE_HEADER_FONT = RGBColor(255, 255, 255)
ALT_ROW_FILL = "F2F6FC"
RISK_COLORS = {
    "High": RGBColor(170, 20, 20),
    "Medium-High": RGBColor(200, 80, 20),
    "Medium": RGBColor(180, 120, 0),
    "Low-Medium": RGBColor(100, 140, 0),
    "Low": RGBColor(0, 100, 0),
}


# ── Helpers ──────────────────────────────────────────────────────
def _font(run, *, bold=None, size_pt=None, color=None, italic=None):
    run.font.name = FONT_BODY
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:ascii"), FONT_BODY)
    rf.set(qn("w:hAnsi"), FONT_BODY)
    rf.set(qn("w:eastAsia"), FONT_BODY)
    run.font.size = Pt(size_pt or FONT_SIZE)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _field(paragraph, instr):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    paragraph._p.append(fld)


def _shade(cell, hex_color):
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc.append(shd)


def _border_bottom(paragraph, color_hex="C0C0C0", sz="6"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), sz)
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), color_hex)
    pbdr.append(b)


def _tune(p, *, justify=True, indent_cm=0.0):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.4
    pf.space_after = Pt(4)
    if indent_cm:
        pf.first_line_indent = Cm(indent_cm)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        _font(r)


def _add_para(doc, text, *, bold=False, color=None, size_pt=None, justify=True, indent_cm=0.0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _font(r, bold=bold, color=color, size_pt=size_pt)
    _tune(p, justify=justify, indent_cm=indent_cm)
    return p


def _add_bullet(doc, text, *, bold_prefix="", indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        _font(rb, bold=True)
    r = p.add_run(text)
    _font(r)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.4
    pf.space_after = Pt(3)
    if indent_level:
        pf.left_indent = Cm(1.27 * indent_level)
    return p


def _add_table(doc, headers, rows, *, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.autofit = True
    for c, h in enumerate(headers):
        cell = t.rows[0].cells[c]
        _shade(cell, TABLE_HEADER_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cp = cell.paragraphs[0]
        cp.text = ""
        cr = cp.add_run(h)
        _font(cr, bold=True, color=TABLE_HEADER_FONT, size_pt=10)
        _tune(cp, justify=False)
    for ri, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = t.rows[ri].cells[c]
            if ri % 2 == 0:
                _shade(cell, ALT_ROW_FILL)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cp = cell.paragraphs[0]
            cp.text = ""
            cr = cp.add_run(str(val))
            color = RISK_COLORS.get(str(val))
            _font(cr, bold=bool(color), color=color, size_pt=10)
            _tune(cp, justify=False)
    doc.add_paragraph()
    return t


# ── Document setup ───────────────────────────────────────────────
def setup(doc: Document):
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(FONT_SIZE)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.4
    normal.paragraph_format.space_after = Pt(4)

    for name, sz, bef, aft in [
        ("Heading 1", 14, 18, 8),
        ("Heading 2", 12, 14, 6),
        ("Heading 3", 11, 10, 5),
    ]:
        s = doc.styles[name]
        s.font.name = FONT_BODY
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        s.font.size = Pt(sz)
        s.font.bold = True
        s.font.color.rgb = COLOR_TITLE
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        s.paragraph_format.line_spacing = 1.4
        s.paragraph_format.space_before = Pt(bef)
        s.paragraph_format.space_after = Pt(aft)
        s.paragraph_format.keep_with_next = True

    for sn in ("List Bullet", "List Number"):
        ls = doc.styles[sn]
        ls.font.name = FONT_BODY
        ls._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        ls.font.size = Pt(FONT_SIZE)

    # Header / footer
    header = sec.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("LEGAL MEMORANDUM | PRIVILEGED & CONFIDENTIAL")
    _font(hr, bold=True, size_pt=8, color=COLOR_MUTED)

    footer = sec.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(fp.add_run("Page "), size_pt=8, color=COLOR_MUTED)
    _field(fp, "PAGE")
    _font(fp.add_run(" of "), size_pt=8, color=COLOR_MUTED)
    _field(fp, "NUMPAGES")


# ── Cover page ───────────────────────────────────────────────────
def cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(p.add_run("PRIVILEGED & CONFIDENTIAL"), bold=True, size_pt=10, color=COLOR_ALERT)

    for _ in range(3):
        doc.add_paragraph()

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(tp.add_run("LEGAL MEMORANDUM"), bold=True, size_pt=18, color=COLOR_TITLE)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(sp.add_run("Comprehensive Legal Risk Review"), bold=True, size_pt=14, color=COLOR_TITLE)

    sp2 = doc.add_paragraph()
    sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(sp2.add_run("Children's Content Engagement Platform"), size_pt=13, color=COLOR_MUTED)
    _border_bottom(sp2, color_hex="19376D", sz="8")

    doc.add_paragraph()
    info = doc.add_table(rows=6, cols=2)
    info.style = "Table Grid"
    info.autofit = True
    meta_rows = [
        ("Prepared by", "Jinju Legal Orchestrator"),
        ("Research Specialist", "Kim Jaesik (김재식)"),
        ("As-of Date", "March 26, 2026"),
        ("Jurisdictions", "Republic of Korea, United States, European Union"),
        ("Expansion Markets", "Japan, Vietnam, Indonesia (overview)"),
        ("Classification", "Privileged & Confidential — Internal Legal Workflow Draft"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        lc = info.rows[i].cells[0]
        rc = info.rows[i].cells[1]
        _shade(lc, "E7EEF8")
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = lc.paragraphs[0]
        lp.text = ""
        _font(lp.add_run(k), bold=True, size_pt=10)
        _tune(lp, justify=False)
        rp = rc.paragraphs[0]
        rp.text = ""
        _font(rp.add_run(v), size_pt=10)
        _tune(rp, justify=False)

    doc.add_paragraph()
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _font(
        dp.add_run(
            "Disclaimer: This memorandum is based on publicly available legislation, regulations, "
            "case law, and industry terms of service as of the date indicated above. It constitutes "
            "general legal research and analysis, not legal advice applicable to any specific "
            "transaction or dispute. The client should consult qualified counsel in each relevant "
            "jurisdiction before making business or legal decisions."
        ),
        size_pt=9,
        color=COLOR_MUTED,
        italic=True,
    )
    doc.add_page_break()


# ── Table of contents ────────────────────────────────────────────
def toc(doc: Document):
    h = doc.add_heading("Table of Contents", level=1)
    for r in h.runs:
        _font(r, bold=True, size_pt=14, color=COLOR_TITLE)
    p = doc.add_paragraph()
    _field(p, r'TOC \o "1-3" \h \z \u')
    _tune(p, justify=False)
    note = doc.add_paragraph(
        "Note: If the table of contents appears blank, open this document in Microsoft Word "
        "and press Ctrl+A then F9 to update all fields."
    )
    _tune(note, justify=False)
    for r in note.runs:
        _font(r, size_pt=9, color=COLOR_MUTED, italic=True)
    doc.add_page_break()


# ── Section builders ─────────────────────────────────────────────
def _risk_tag(doc, level):
    p = doc.add_paragraph()
    r = p.add_run(f"Risk Level: {level}")
    _font(r, bold=True, color=RISK_COLORS.get(level, COLOR_MUTED), size_pt=FONT_SIZE)
    _tune(p, justify=False)


def _section_head(doc, num, title, level=1):
    h = doc.add_heading(f"{num}. {title}", level=level)
    for r in h.runs:
        _font(r, bold=True, size_pt=[0, 14, 12, 11][level], color=COLOR_TITLE)


def _sub_head(doc, num, title):
    h = doc.add_heading(f"{num} {title}", level=2)
    for r in h.runs:
        _font(r, bold=True, size_pt=12, color=COLOR_TITLE)


def _sub_sub_head(doc, num, title):
    h = doc.add_heading(f"{num} {title}", level=3)
    for r in h.runs:
        _font(r, bold=True, size_pt=11, color=COLOR_TITLE)


# ── Main content ─────────────────────────────────────────────────
def executive_summary(doc):
    _section_head(doc, "I", "Executive Summary")
    _add_para(
        doc,
        "This memorandum presents a comprehensive legal risk review of the proposed Children's "
        "Content Engagement Platform—an AI-powered service that analyzes children's television "
        "shows, movies, and videos at the episode level and generates structured 'engagement "
        "packs' (discussion questions, activity guides, and learning extensions) for parents. "
        "The platform plans to launch in the Korean market before expanding to the United States, "
        "the European Union, Japan, Vietnam, and Indonesia."
    )
    _add_para(
        doc,
        "We have analyzed the business model across seven primary risk domains: (1) copyright "
        "and intellectual property; (2) OTT platform terms-of-service compliance; (3) privacy "
        "and children's data regulation; (4) AI-related regulation; (5) child-facing service "
        "regulations; (6) B2B model risks; and (7) other regulatory and commercial risks "
        "including endorsement disclosure, expansion-market regulation, and government-grant "
        "conditions."
    )

    _add_para(doc, "Overall Assessment:", bold=True)
    _add_para(
        doc,
        "The business model is legally viable, but faces several areas of elevated risk that "
        "require proactive mitigation before launch. The highest-priority risks are: "
        "(a) OTT platform terms-of-service violations from automated content analysis; "
        "(b) copyright exposure from the use of screenshots or frame captures; and "
        "(c) children's data compliance obligations under Korea's PIPA, the U.S. COPPA, and "
        "the EU GDPR. Each of these risks is manageable through the design choices and "
        "operational measures recommended below."
    )

    _add_table(
        doc,
        ["Risk Domain", "Risk Level", "Priority Action"],
        [
            ["Copyright (AI analysis & summaries)", "Medium", "Avoid verbatim reproduction; pursue studio licenses"],
            ["Copyright (screenshots / frames)", "High", "Do not use frame captures without license"],
            ["OTT Terms of Service", "High", "Use legitimate viewing + AI vision; pursue partnerships"],
            ["Children's Data (KR PIPA)", "High", "Implement parental consent system for under-14"],
            ["Children's Data (US COPPA)", "Medium-High", "Design as parent-directed; minimize child data"],
            ["Children's Data (EU GDPR/AADC)", "Medium-High", "DPIA required; comply with Age Appropriate Design Code"],
            ["AI Regulation (EU AI Act)", "Medium", "Implement transparency disclosures by Aug 2026"],
            ["Advertising / Sponsorship Disclosure", "Medium-High", "Editorial independence charter; clear labels"],
            ["Gamification (children)", "Medium", "Avoid manipulative nudge patterns"],
            ["B2B Education Compliance", "Low-Medium", "Standard data processing agreements"],
            ["Expansion Markets", "Medium", "Jurisdiction-specific legal review before entry"],
            ["Government Grants (TIPS)", "Low-Medium", "Comply with IP/milestone reporting conditions"],
        ],
    )
    doc.add_page_break()


def section_copyright(doc):
    _section_head(doc, "II", "Copyright Risk")

    # ── 1. AI Analysis & Fair Use ──
    _sub_head(doc, "A.", "AI-Based Content Analysis and Fair Use")

    _add_para(doc, "Risk Summary:", bold=True)
    _add_para(
        doc,
        "The platform's core function—using AI to watch copyrighted children's content and "
        "generate original discussion guides, educational themes, and activity suggestions—"
        "engages copyright law in every target jurisdiction. The critical question is whether "
        "this activity constitutes permissible fair use (or its equivalent) or whether it "
        "requires a license from the copyright holder."
    )

    _sub_sub_head(doc, "(i)", "Korea — Copyright Act Article 35-5 (Fair Use)")
    _add_para(
        doc,
        "Korea's Copyright Act (저작권법) Article 35-5 provides a general fair-use provision "
        "modeled on the U.S. four-factor test. It permits use of copyrighted works where such "
        "use (1) does not conflict with normal exploitation of the work and (2) does not "
        "unreasonably prejudice the legitimate interests of the author. The statute directs "
        "courts to consider four factors:"
    )
    _add_bullet(doc, "The purpose and character of the use;", bold_prefix="Factor 1: ")
    _add_bullet(doc, "The type and intended purpose of the copyrighted work;", bold_prefix="Factor 2: ")
    _add_bullet(doc, "The portion of the work used in relation to the whole and its significance;", bold_prefix="Factor 3: ")
    _add_bullet(doc, "The effect of the use on the current or potential market value of the work.", bold_prefix="Factor 4: ")

    _add_para(
        doc,
        "Applied to this platform, Factors 1 and 4 favor the platform. The AI does not reproduce "
        "the content itself but generates entirely new, transformative works (discussion questions, "
        "educational guides) that serve a different purpose—parent-child engagement—than the "
        "original entertainment purpose. The platform's output does not substitute for watching "
        "the show; indeed, it is designed to encourage viewership. Factor 2 is neutral (children's "
        "shows are creative works, which weighs slightly against fair use, but they are also "
        "commercially published). Factor 3 depends on implementation: if the platform does not "
        "reproduce substantial portions of dialogue, plot summaries detailed enough to substitute "
        "for viewing, or visual content, this factor favors the platform."
    )
    _add_para(
        doc,
        "Korea's fair-use provision is relatively new (introduced in 2011) with limited case law. "
        "Courts have generally construed it narrowly. There is no Korean precedent directly on "
        "point for AI-based content analysis. However, Article 28 (quotation for criticism, "
        "education, or research) provides additional support for quoting limited excerpts in "
        "an educational context, provided the quotation is within 'fair scope' and 'consistent "
        "with fair practices.'",
    )
    _add_para(
        doc,
        "Penalties: Copyright infringement under Article 136(1) carries up to 5 years' "
        "imprisonment or KRW 50 million (approx. USD 37,000) in criminal fines, plus civil "
        "damages. Corporate liability applies under Article 141.",
    )

    _sub_sub_head(doc, "(ii)", "United States — 17 U.S.C. § 107 (Fair Use)")
    _add_para(
        doc,
        "U.S. fair-use doctrine provides a stronger foundation for this business model. The "
        "leading precedent is Authors Guild v. Google, Inc., 804 F.3d 202 (2d Cir. 2015), "
        "where the Second Circuit held that Google's digitization of entire library books "
        "and provision of search results with limited snippets was 'highly transformative' fair "
        "use. The court emphasized that the use served a different purpose (information discovery) "
        "from the original (reading) and did not substitute for the original work."
    )
    _add_para(
        doc,
        "However, the Supreme Court's decision in Andy Warhol Foundation for the Visual Arts v. "
        "Goldsmith, 598 U.S. 508 (2023), narrowed the transformative-use analysis. The Court "
        "held that where the use has the same or highly similar purpose as the original, the "
        "first fair-use factor weighs against the user regardless of how much the work has been "
        "altered. For this platform, the distinction is favorable: discussion guides and "
        "educational engagement packs serve a fundamentally different purpose (parental engagement "
        "and education) than the original (children's entertainment). This positions the platform "
        "closer to Google Books than to the Warhol scenario."
    )
    _add_para(
        doc,
        "More recent decisions provide both supporting and cautionary precedent. In Bartz v. "
        "Anthropic PBC, No. 3:23-cv-07721 (N.D. Cal. June 23, 2025), and Kadrey v. Meta "
        "Platforms, Inc., No. 3:23-cv-03417 (N.D. Cal. June 25, 2025), courts held that using "
        "copyrighted books to train large language models was 'quintessentially transformative' "
        "fair use because the purpose (training a general-purpose AI model) is fundamentally "
        "different from the purpose of the original works. These rulings support the proposition "
        "that AI analysis of copyrighted content for a different purpose may qualify as fair use."
    )
    _add_para(
        doc,
        "Conversely, Thomson Reuters Enterprise Centre GmbH v. Ross Intelligence Inc., No. "
        "1:20-cv-613 (D. Del. Feb. 11, 2025), was the first court to reject fair use for AI "
        "training on copyrighted text. The court found that Ross Intelligence's AI legal research "
        "tool was a direct 'market substitute' for Westlaw (Thomson Reuters' product), and that "
        "the fourth factor—market effect—was decisive. This precedent warns that if the platform's "
        "discussion guides compete with officially licensed companion materials or educational "
        "products, the fair-use defense weakens significantly."
    )
    _add_para(
        doc,
        "Also instructive is Warner Bros. Entertainment Inc. v. RDR Books, 575 F. Supp. 2d 513 "
        "(S.D.N.Y. 2008), where a Harry Potter reference guide was held to infringe because it "
        "reproduced too much original expression without sufficient transformative commentary. "
        "This underscores the importance of the platform generating genuinely original analysis "
        "rather than detailed plot recapitulations."
    )
    _add_para(
        doc,
        "An additional critical concern is DMCA Section 1201 (17 U.S.C. § 1201), which prohibits "
        "circumventing technological protection measures (TPMs) that control access to copyrighted "
        "works. Streaming services employ DRM systems (Widevine, FairPlay) that are TPMs under "
        "Section 1201. Even for legitimate subscribers, circumventing DRM to capture video frames "
        "for AI analysis likely violates Section 1201(a)(1)(A), independently of any fair-use "
        "analysis. No Section 1201 exemption currently exists for AI content analysis of streaming "
        "media. Recent cases including Reddit v. Perplexity AI (2025) and NYT v. Perplexity AI "
        "(Dec. 2025) have advanced anti-circumvention theories against AI services that bypass "
        "access controls."
    )

    _sub_sub_head(doc, "(iii)", "European Union — Copyright Directive 2019/790")
    _add_para(
        doc,
        "The EU Copyright Directive addresses text and data mining (TDM) in Articles 3 and 4. "
        "Article 3 permits TDM by research organizations and cultural heritage institutions "
        "with lawful access. Article 4 permits TDM by any person with lawful access, but "
        "rights holders may expressly 'opt out' by reserving their rights in machine-readable "
        "form."
    )
    _add_para(
        doc,
        "A commercial platform that systematically analyzes copyrighted content likely falls "
        "under Article 4, not Article 3. If Netflix, Disney, or other rights holders opt out "
        "of TDM (and their terms of service may constitute such a reservation), the platform "
        "would need to obtain licenses for EU-market operations. This represents a meaningful "
        "constraint on the EU expansion path."
    )
    _add_para(
        doc,
        "Additionally, the act of sampling video frames may constitute 'reproduction' under "
        "Article 2 of the InfoSoc Directive (2001/29/EC), requiring either a license or an "
        "applicable exception."
    )

    _risk_tag(doc, "Medium")
    _add_para(doc, "Recommended Mitigation Measures:", bold=True)
    _add_bullet(doc, "Do not reproduce verbatim dialogue, scripts, or detailed plot summaries that could substitute for viewing the content.")
    _add_bullet(doc, "Generate original, transformative analysis (discussion questions, developmental themes, activity suggestions) rather than derivative summaries.")
    _add_bullet(doc, "Maintain clear editorial separation between AI-generated analysis and any referenced content.")
    _add_bullet(doc, "Pursue content partnership or licensing agreements with studios as the business scales.")
    _add_bullet(doc, "Monitor the TDM opt-out status of rights holders for EU market entry.")

    # ── 2. Subtitles & Scripts ──
    _sub_head(doc, "B.", "Copyrightability of Subtitles and Scripts")
    _add_para(doc, "Risk Summary:", bold=True)
    _add_para(
        doc,
        "Official subtitles and scripts are copyrightable literary works under Korean, U.S., and "
        "EU law. The use of official subtitle files—whether obtained from streaming platforms, "
        "OpenSubtitles, or other sources—without authorization constitutes copyright infringement. "
        "OpenSubtitles explicitly prohibits commercial use of its database."
    )
    _add_para(
        doc,
        "AI-generated transcriptions from the audio track produce independent works through the "
        "platform's own effort, but may contain expression substantially similar to the official "
        "scripts, particularly for simple dialogue. The transcription itself is lawful (the "
        "platform is processing audio it lawfully received as a subscriber), but publishing or "
        "distributing verbatim transcriptions commercially may infringe the underlying script's "
        "copyright."
    )
    _risk_tag(doc, "High")
    _add_para(doc, "Recommended Mitigation Measures:", bold=True)
    _add_bullet(doc, "Never access, copy, or distribute official subtitle files from any source.")
    _add_bullet(doc, "Use AI audio analysis only for internal content understanding; do not publish verbatim transcriptions.")
    _add_bullet(doc, "Focus outputs on original analysis (themes, questions, activities) rather than on reproducing dialogue.")

    # ── 3. Screenshots / Frame Captures ──
    _sub_head(doc, "C.", "Screenshots and Frame Captures")
    _add_para(doc, "Risk Summary:", bold=True)
    _add_para(
        doc,
        "Individual frames from copyrighted audiovisual works are protectable as component parts "
        "of the larger work. Using screenshots or frame captures in commercial engagement packs "
        "without authorization creates direct infringement risk. While some U.S. courts have "
        "permitted thumbnail reproductions in search contexts—Kelly v. Arriba Soft Corp., 336 "
        "F.3d 811 (9th Cir. 2003) (thumbnails in image search = fair use); Perfect 10, Inc. v. "
        "Amazon.com, Inc., 508 F.3d 1146 (9th Cir. 2007) (inline linking distinguished from "
        "direct copying)—these precedents involve non-commercial search functionality, not "
        "commercial educational products."
    )
    _add_para(
        doc,
        "Korean law provides no equivalent thumbnail exception. Using frame captures from "
        "children's shows in paid engagement packs or subscription content would likely exceed "
        "the scope of Article 35-5 fair use, particularly given the commercial purpose and the "
        "availability of licensing alternatives."
    )
    _add_para(
        doc,
        "Additionally, publicity and personality rights may be implicated where recognizable "
        "animated characters appear in frame captures. Character merchandising rights in Korea "
        "are protected under trademark law, unfair competition law, and civil personality rights."
    )

    _risk_tag(doc, "High")
    _add_para(doc, "Recommended Mitigation Measures:", bold=True)
    _add_bullet(doc, "Do not capture, store, or display screenshots or frame captures from copyrighted content in any user-facing output.")
    _add_bullet(doc, "Use text-based descriptions, original illustrations, or AI-generated imagery that does not reproduce copyrighted character designs.")
    _add_bullet(doc, "If visual content is essential, use only officially licensed press/promotional images under a written agreement.")
    _add_bullet(doc, "For the content creator dashboard, use aggregated data visualizations rather than actual content frames.")

    doc.add_page_break()


def section_tos(doc):
    _section_head(doc, "III", "OTT Platform Terms of Service Violations")

    _add_para(doc, "Risk Summary:", bold=True)
    _add_para(
        doc,
        "The major streaming platforms expressly prohibit commercial analysis, automated "
        "processing, and derivative use of their content in their subscriber agreements. The "
        "platform's business model—which depends on AI analysis of content available on Netflix, "
        "Disney+, YouTube, and other services—creates a direct tension with these contractual "
        "restrictions."
    )

    _sub_head(doc, "A.", "Netflix Terms of Use")
    _add_para(
        doc,
        "Netflix's Terms of Use broadly restrict subscriber rights. Relevant provisions include "
        "prohibitions on: copying, reproducing, or redistributing content; using automated "
        "systems to access the service; and any commercial exploitation of Netflix content or "
        "information obtained through the service. Netflix shut down its public API in November "
        "2014, and its current terms explicitly prohibit use of content for AI or machine-"
        "learning purposes."
    )

    _sub_head(doc, "B.", "Disney+ Subscriber Agreement")
    _add_para(
        doc,
        "The Disney+ Subscriber Agreement contains the most explicit anti-AI provisions among "
        "major streaming platforms. Users may not access, monitor, copy, or extract the Services "
        "'using a robot, spider, script or other automated means, including for purposes of "
        "creating or developing any AI Tool, data mining or web scraping or otherwise compiling, "
        "building, creating or contributing to any collection of data, data set or database.' "
        "This language directly targets the type of AI-powered content analysis contemplated by "
        "this platform. Disney+ also limits individual subscriber liability to $1,000 and "
        "requires binding individual arbitration with class action waivers."
    )

    _sub_head(doc, "C.", "YouTube Terms of Service and API Terms")
    _add_para(
        doc,
        "YouTube's Terms of Service prohibit automated access and commercial exploitation of "
        "content. Section 5 restricts: accessing the service using automated means (bots, "
        "scrapers); collecting personal information; and reproducing, distributing, or creating "
        "derivative works from content. YouTube's API Terms of Service impose additional "
        "restrictions on data usage, display, and caching. However, YouTube's ecosystem is "
        "somewhat more permissive than closed platforms, as some creators may welcome analysis "
        "and engagement tools."
    )

    _sub_head(doc, "D.", "Legal Consequences and Enforceability")
    _add_para(
        doc,
        "Terms-of-service violations expose the platform to several categories of legal risk:"
    )
    _add_bullet(doc, "Account termination and platform bans (immediate and automatic).", bold_prefix="Contract remedies: ")
    _add_bullet(doc, "Breach of contract claims for compensatory damages (typically difficult to quantify for ToS breaches).", bold_prefix="Damages: ")
    _add_bullet(doc, "Following Van Buren v. United States, 593 U.S. 374 (2021), the U.S. Supreme Court held that exceeding authorized access under the CFAA requires accessing information to which the user was not entitled at all, not merely violating usage restrictions. Therefore, CFAA liability is unlikely for ToS violations by authorized subscribers.", bold_prefix="CFAA (U.S.): ")
    _add_bullet(doc, "Korea's 정보통신망법 (Network Act) does not criminalize ToS violations per se, but unauthorized access to computer systems may implicate criminal provisions.", bold_prefix="Korean criminal law: ")
    _add_bullet(doc, "Platform operators could allege tortious interference with business relationships or unfair competition in tort.", bold_prefix="Tort theories: ")

    _add_para(
        doc,
        "The enforceability of ToS provisions that purport to prohibit all analysis of content "
        "is debatable. In MDY Industries, LLC v. Blizzard Entertainment, Inc. (9th Cir. 2010), "
        "the Ninth Circuit held that a ToS violation constitutes copyright infringement only "
        "where there is a 'nexus between the condition and the licensor's exclusive rights of "
        "copyright.' However, MDY was liable under DMCA Section 1201 for circumvention. "
        "Whether contractual restrictions can override copyright fair-use exceptions remains "
        "an unsettled question with a deepening circuit split. Most courts have found the "
        "Copyright Act does not preempt contractual constraints, meaning ToS restrictions that "
        "exceed copyright law are enforceable as contracts even if they limit uses that would "
        "otherwise constitute fair use."
    )

    _risk_tag(doc, "High")
    _add_para(doc, "Recommended Mitigation Measures:", bold=True)
    _add_bullet(doc, "Use only legitimate subscriber access; view content through normal playback rather than automated scraping or API exploitation.")
    _add_bullet(doc, "Deploy AI vision analysis (frame sampling during lawful viewing) rather than accessing backend data, subtitle files, or APIs.")
    _add_bullet(doc, "Prioritize studio partnership agreements that explicitly authorize content analysis, beginning with Korean studios (e.g., ICONIX/Pororo) that have aligned incentives.")
    _add_bullet(doc, "Bootstrap initial content from public-domain and Creative Commons-licensed sources (Wikipedia, fan wikis, public educational materials) to reduce dependence on ToS-restricted platforms.")
    _add_bullet(doc, "Maintain a clear legal position that discussion guides and educational analysis are independently created transformative works, not derivatives of the underlying content.")
    _add_bullet(doc, "Consult with litigation counsel before scaling AI analysis of content from platforms with explicit AI/ML prohibitions.")

    doc.add_page_break()


def section_privacy(doc):
    _section_head(doc, "IV", "Privacy and Children's Data Regulation")

    _add_para(doc, "Risk Summary:", bold=True)
    _add_para(
        doc,
        "Children's data protection is among the most heavily regulated areas in all three "
        "target jurisdictions. Because this platform serves families with minor children and may "
        "collect data about children's viewing habits, age-based preferences, and developmental "
        "information, compliance with children's data regulations is a non-negotiable prerequisite "
        "for launch."
    )

    _sub_head(doc, "A.", "Korea — Personal Information Protection Act (PIPA)")
    _add_para(
        doc,
        "Korea's PIPA (개인정보 보호법), as amended through 2025, imposes specific obligations "
        "for processing personal information of children under 14. Article 22-2 provides:"
    )
    _add_bullet(doc, "To process personal information of a child under 14, the controller must obtain the consent of the child's legal guardian (법정대리인).")
    _add_bullet(doc, "The controller must verify that the guardian has actually consented (not merely claim consent was obtained).")
    _add_bullet(doc, "Notices to children under 14 must be in easy-to-understand format with clear, simple language.")
    _add_bullet(doc, "The PIPA Enforcement Decree specifies permissible methods for guardian consent verification, including online name/identity verification, credit card authentication, and mobile phone authentication.")

    _add_para(
        doc,
        "Penalties for non-compliance are severe. Article 71(3) provides that processing a "
        "child's personal information without guardian consent is punishable by up to 5 years' "
        "imprisonment or a fine of up to KRW 50 million (approximately USD 37,000). Additionally, "
        "the Personal Information Protection Commission (PIPC) may impose administrative fines "
        "(과징금) of up to 3% of relevant revenue and corrective orders."
    )
    _add_para(
        doc,
        "Even though the platform is designed for parents, if it collects any information "
        "about the child (e.g., age, name, viewing history, interests for personalized "
        "engagement packs), PIPA Article 22-2 is triggered. The platform must implement a "
        "verifiable parental consent mechanism before any child-related data is collected."
    )

    _sub_head(doc, "B.", "United States — COPPA")
    _add_para(
        doc,
        "The Children's Online Privacy Protection Act (COPPA), 15 U.S.C. §§ 6501–6506, and "
        "the FTC's implementing Rule (16 C.F.R. Part 312) apply to (a) websites and online "
        "services 'directed to children' under 13, and (b) general-audience services that have "
        "'actual knowledge' they are collecting personal information from children under 13."
    )
    _add_para(
        doc,
        "Classification of this platform requires careful analysis. The platform is directed "
        "to parents, not to children directly. Parents create accounts and manage the service. "
        "Under the FTC's guidance, a service directed to parents (not children) that collects "
        "information about children is not automatically 'directed to children' for COPPA "
        "purposes. However, if the platform's content, advertising, or design features suggest "
        "it targets children as primary users (e.g., gamification features, child-facing UI "
        "elements), the FTC may classify it as a 'mixed audience' service subject to COPPA."
    )
    _add_para(
        doc,
        "The FTC finalized significant COPPA Rule amendments on January 16, 2025, effective "
        "June 23, 2025, with a compliance deadline of April 22, 2026. Key changes include: "
        "(a) separate opt-in consent required before using children's data for targeted "
        "advertising; (b) expanded definition of 'personal information' to include biometric "
        "identifiers; (c) new explicit data retention limits; and (d) new verifiable parental "
        "consent methods including knowledge-based authentication and facial recognition with "
        "immediate deletion. Penalties for COPPA violations can reach $53,088 per violation per "
        "day. Recent enforcement actions include HoYoverse/Genshin Impact ($20 million, 2024), "
        "Walt Disney ($10 million, 2024), and NGL Labs ($5 million, 2024)."
    )

    _sub_head(doc, "C.", "EU — GDPR and Age Appropriate Design Code")
    _add_para(
        doc,
        "The GDPR imposes heightened protections for children's data. Article 8 provides that "
        "where consent is relied upon for offering 'information society services' directly to a "
        "child, the child's consent is valid only if the child is at least 16 years old (or a "
        "lower age as set by member states, with a floor of 13). Below that threshold, the "
        "controller must obtain consent from the holder of parental responsibility."
    )
    _add_para(
        doc,
        "Article 6(1)(f) (legitimate interest) may not be relied upon for children's data "
        "processing where the controller's interest is overridden by the child's fundamental "
        "rights and freedoms. Recital 38 emphasizes that children 'merit specific protection' "
        "regarding their personal data, 'in particular with regard to the use of [their data] "
        "for the purposes of marketing or creating personality or user profiles.'"
    )
    _add_para(
        doc,
        "The UK's Age Appropriate Design Code (AADC, also known as the Children's Code) "
        "establishes 15 standards for online services likely to be accessed by children under "
        "18. Key requirements include: 'best interests of the child' as a primary design "
        "consideration; data minimization by default; prohibition of nudge techniques that "
        "encourage children to provide unnecessary personal data; privacy settings set to "
        "'high privacy' by default; and a Data Protection Impact Assessment (DPIA) for services "
        "likely to be accessed by children."
    )
    _add_para(
        doc,
        "AADC enforcement has been aggressive: TikTok was fined GBP 12.7 million in 2023 for "
        "failing to use children's data lawfully, and Reddit was fined GBP 14.47 million in "
        "February 2026 for unlawful processing of children's personal information, including "
        "failures in age assurance and DPIA compliance. Penalties can reach GBP 17.5 million or "
        "4% of global annual turnover. While the AADC is UK-specific, the EDPB has issued "
        "similar guidance, and a DPIA is mandatory before launching any feature that processes "
        "children's data in the EU market."
    )

    _risk_tag(doc, "High")
    _add_para(doc, "Recommended Mitigation Measures:", bold=True)
    _add_bullet(doc, "Korea: Implement verifiable parental consent (e.g., mobile phone or identity verification) for any data collection related to children under 14 before Korean market launch.")
    _add_bullet(doc, "U.S.: Design the platform as parent-directed (parents create and manage all accounts). Minimize collection of child-specific data. If child data is collected, implement COPPA-compliant verifiable parental consent.")
    _add_bullet(doc, "EU: Conduct a DPIA before launch. Default to high-privacy settings. Implement age verification and parental consent mechanisms. Comply with AADC standards for any feature accessible to children.")
    _add_bullet(doc, "All jurisdictions: Adopt a dedicated children's privacy policy; implement data minimization and retention limits; provide easily accessible deletion mechanisms.")

    doc.add_page_break()


def section_ai_regulation(doc):
    _section_head(doc, "V", "AI-Related Regulation")

    _sub_head(doc, "A.", "EU AI Act — Regulation (EU) 2024/1689")
    _add_para(doc, "Risk Summary:", bold=True)
    _add_para(
        doc,
        "The EU AI Act, which entered into force on August 1, 2024, establishes a risk-based "
        "regulatory framework for AI systems. The platform's AI components—content analysis, "
        "review generation, age-calibrated recommendation engine, and the 'What Should We Watch "
        "Tonight?' chatbot—must be assessed under this framework."
    )
    _add_para(
        doc,
        "Classification analysis: Annex III of the AI Act lists 'high-risk' use cases, "
        "including Category 3 (education and vocational training—AI systems used to determine "
        "access to or performance in educational institutions) and Category 4 (employment). "
        "The platform's engagement packs are educational in nature, but they do not determine "
        "access to or performance in educational institutions. Therefore, the platform is more "
        "likely classified as a 'limited risk' system, subject to transparency obligations "
        "under Article 50 rather than the full high-risk compliance framework."
    )
    _add_para(
        doc,
        "Critically, Article 5(1)(b) of the AI Act—which has been in force since February 2, "
        "2025—prohibits AI systems that exploit vulnerabilities of specific groups due to their "
        "age (including children) to materially distort behavior in a manner likely to cause "
        "significant harm. If the recommendation engine uses manipulative techniques exploiting "
        "children's cognitive vulnerabilities to drive engagement, it could trigger this outright "
        "prohibition, with penalties up to EUR 35 million or 7% of global annual turnover."
    )
    _add_para(
        doc,
        "Article 50 transparency requirements for limited-risk AI systems include: disclosing "
        "to users that content (reviews, recommendations, discussion questions) has been "
        "AI-generated; ensuring that AI-generated content is marked in machine-readable format; "
        "and informing users interacting with the chatbot that they are communicating with an "
        "AI system. These obligations apply from August 2, 2026. The European Commission "
        "published a first draft Code of Practice on transparency of AI-generated content in "
        "December 2025, with a final code expected June 2026."
    )
    _add_para(
        doc,
        "For the recommendation chatbot ('What Should We Watch Tonight?'), Article 50(1) "
        "requires that the user be informed that they are interacting with an AI system, unless "
        "this is obvious from the circumstances."
    )

    _sub_head(doc, "B.", "AI Content Reliability and Platform Liability")
    _add_para(
        doc,
        "The platform faces reputational and potentially legal liability if AI-generated content "
        "reviews or age ratings contain errors, biases, or misleading assessments. Specific risks "
        "include:"
    )
    _add_bullet(doc, "An AI system that incorrectly rates violent or age-inappropriate content as suitable for young children could cause harm and generate liability under general negligence principles or consumer protection law.", bold_prefix="False negatives: ")
    _add_bullet(doc, "AI analysis may reflect biases present in training data, including cultural, racial, or gender stereotypes. If engagement packs perpetuate harmful stereotypes, the platform could face regulatory scrutiny and reputational damage.", bold_prefix="AI bias: ")
    _add_bullet(doc, "Korea currently lacks comprehensive AI-specific legislation, but the government has published 'AI Ethics Standards' (인공지능 윤리기준, 2020) and is developing binding legislation. The National AI Committee is preparing a 'Basic Act on AI' with provisions relevant to high-impact AI services.", bold_prefix="Korean AI regulation: ")

    _risk_tag(doc, "Medium")
    _add_para(doc, "Recommended Mitigation Measures:", bold=True)
    _add_bullet(doc, "Clearly label all AI-generated content as such (e.g., 'AI-Generated Review' or 'AI-Suggested Discussion Questions').")
    _add_bullet(doc, "Implement human review for age-rating assignments, especially for content that may involve violence, sexual themes, or other sensitive material.")
    _add_bullet(doc, "Establish a user feedback and correction mechanism for AI errors.")
    _add_bullet(doc, "Conduct regular bias audits of AI output, particularly for cultural sensitivity in the Korean market.")
    _add_bullet(doc, "Prepare for EU AI Act transparency compliance by August 2026.")

    doc.add_page_break()


def section_child_services(doc):
    _section_head(doc, "VI", "Regulations Specific to Child-Facing Services")

    _sub_head(doc, "A.", "Korea — Youth Protection Act and Network Act")
    _add_para(
        doc,
        "The Youth Protection Act (청소년 보호법) defines 'youth' as persons under 19 and "
        "establishes a framework for regulating 'youth harmful media' (청소년유해매체물). "
        "The platform itself is not 'harmful media' — it is an educational supplement for "
        "parents. However, if the platform provides content (such as detailed descriptions of "
        "violent or sexual scenes in children's shows, even for analytical purposes), it must "
        "ensure compliance with the Act's content standards."
    )
    _add_para(
        doc,
        "Article 9 of the Youth Protection Act lists criteria for harmful media "
        "determinations, including content that: stimulates sexual desire (Art. 9(1)(i)); "
        "could provoke violence or criminal impulses (Art. 9(1)(ii)); glorifies violence or "
        "substance abuse (Art. 9(1)(iii)); or promotes gambling (Art. 9(1)(iv))."
    )
    _add_para(
        doc,
        "The Network Act (정보통신망 이용촉진 및 정보보호 등에 관한 법률) imposes additional "
        "obligations on information service providers. Article 42-3 requires designation of a "
        "youth protection officer (청소년 보호 책임자) for service providers meeting certain "
        "daily-user and revenue thresholds. Article 42 requires providers offering youth harmful "
        "media to apply appropriate labeling and access restrictions."
    )
    _risk_tag(doc, "Low-Medium")

    _sub_head(doc, "B.", "Advertising and Sponsorship Regulations for Children")
    _add_para(
        doc,
        "The studio sponsorship revenue model—where content studios pay for 'featured placement, "
        "expedited review processing, and engagement pack creation services'—creates significant "
        "endorsement disclosure obligations."
    )
    _add_para(doc, "Korean Law:", bold=True)
    _add_para(
        doc,
        "The Fair Labeling and Advertising Act (표시·광고의 공정화에 관한 법률) prohibits "
        "deceptive or unfair advertising. The KFTC's Guidelines on Endorsement and Recommendation "
        "Advertising (추천·보증 등에 관한 표시·광고 심사지침), as amended effective December 1, "
        "2024, now require that economic relationships be disclosed at the title or beginning of "
        "the content, not buried in descriptions or hashtags. The amendment broadened 'economic "
        "relationship' (경제적 이해관계) to cover monetary payment, free products, discounts, "
        "early access, conditional rewards, and any benefit affecting objectivity. Where a studio "
        "sponsors the creation of an engagement pack for its own content, this constitutes a "
        "'material connection' requiring upfront disclosure. Both the sponsoring studio and the "
        "platform bear responsibility for disclosure."
    )
    _add_para(doc, "U.S. Law:", bold=True)
    _add_para(
        doc,
        "The FTC Endorsement Guides (16 C.F.R. Part 255, updated June 2023) require clear and "
        "conspicuous disclosure of material connections between endorsers and advertisers. For "
        "content 'directed to children,' the FTC applies heightened scrutiny under its general "
        "authority over unfair and deceptive practices (Section 5 of the FTC Act). The platform's "
        "position that 'studios cannot pay to change review scores' does not eliminate the "
        "disclosure obligation — the material relationship itself must be disclosed."
    )
    _add_para(
        doc,
        "The critical risk is the perception of review bias. If parents perceive that sponsored "
        "content receives favorable treatment (even through 'priority placement' rather than "
        "altered scores), this undermines the platform's core value proposition of trustworthy, "
        "independent content analysis."
    )

    _risk_tag(doc, "Medium-High")
    _add_para(doc, "Recommended Mitigation Measures:", bold=True)
    _add_bullet(doc, "Implement clear, unambiguous 'Sponsored' labels on all studio-sponsored content, visible at the point of consumption.")
    _add_bullet(doc, "Adopt a written editorial independence charter that prohibits studios from influencing review scores, age ratings, or content analysis.")
    _add_bullet(doc, "Physically and organizationally separate the editorial function (content analysis) from the commercial function (sponsorship sales).")
    _add_bullet(doc, "Disclose the sponsorship model transparently in the platform's about/FAQ sections.")
    _add_bullet(doc, "Apply FTC Endorsement Guide standards as the baseline for all markets.")

    _sub_head(doc, "C.", "Gamification Features — Regulatory Risk")
    _add_para(
        doc,
        "The 'Watch Together' badges, streaks, and achievement systems raise regulatory "
        "concerns under children's design standards, particularly the UK Age Appropriate Design "
        "Code (AADC) Standard 10 ('Nudge Reduction'). The AADC prohibits design features that "
        "'nudge, lead or encourage children to provide unnecessary personal data, weaken or turn "
        "off their privacy protections, or extend their use.' Streaks and achievement mechanics "
        "are specifically identified as potentially manipulative design patterns for children."
    )
    _add_para(
        doc,
        "The EU Digital Services Act (Regulation (EU) 2022/2065) Article 25 prohibits 'dark "
        "patterns'—interface designs that materially distort or impair user decision-making. "
        "While the current definition targets deceptive rather than addictive patterns, the "
        "legislative trend is toward broader regulation of manipulative design in children's "
        "services."
    )
    _add_para(
        doc,
        "Under Korean law, the amended Act on Consumer Protection in Electronic Commerce "
        "(전자상거래법), effective February 14, 2025, specifically prohibits six types of dark "
        "patterns: drip pricing, hidden renewals, nagging, pre-selected options, misleading "
        "hierarchy, and obstructing cancellation/withdrawal. While not children-specific, these "
        "prohibitions apply to any online service. Gamification reward loops that constitute "
        "'nagging' (반복적 요청) or use pre-selected engagement options are covered. The KFTC "
        "has already identified 45 suspected cases in its initial monitoring period (February–"
        "July 2025)."
    )
    _risk_tag(doc, "Medium")
    _add_para(doc, "Recommended Mitigation Measures:", bold=True)
    _add_bullet(doc, "Design gamification features to reward parent-child co-viewing, not individual child screen time.")
    _add_bullet(doc, "Avoid loss-framing mechanics (e.g., 'You'll lose your streak!') and limit notification frequency.")
    _add_bullet(doc, "Ensure badge and streak features are accessible to parents, not directly to children.")
    _add_bullet(doc, "Comply with AADC Standard 10 as a design baseline for all markets.")

    doc.add_page_break()


def section_b2b(doc):
    _section_head(doc, "VII", "B2B Model Risks")

    _sub_head(doc, "A.", "Education-Related Compliance for Schools and Hagwons")
    _add_para(
        doc,
        "The B2B revenue stream targeting Korean schools and hagwons (private academies) "
        "introduces education-specific regulatory considerations. The Act on the Establishment "
        "and Operation of Private Teaching Institutes (학원의 설립·운영 및 과외교습에 관한 법률) "
        "regulates hagwon operations but does not directly regulate content providers to "
        "hagwons. The platform is a content/tool supplier, not a hagwon, and is therefore not "
        "subject to direct hagwon registration or operational requirements."
    )
    _add_para(
        doc,
        "A critical distinction: Korean Copyright Act Article 25 permits formal educational "
        "institutions (정규 교육기관 established under special laws) to reproduce and distribute "
        "copyrighted works for classroom use, subject to a compensation payment system. However, "
        "hagwons—as private academies (사교육기관)—are explicitly excluded from this educational "
        "use exception. Therefore, hagwons cannot rely on Article 25 and must secure independent "
        "copyright licenses for all content used in instruction. This makes the B2B content "
        "license agreement the primary legal basis for a hagwon's lawful use of the platform's "
        "materials."
    )
    _add_para(
        doc,
        "When selling to public schools, the platform must consider the Framework Act on "
        "Education (교육기본법) and applicable procurement regulations. Content provided to "
        "schools may need to align with the national curriculum framework, and procurement "
        "may require compliance with government e-procurement procedures (나라장터 / GePS)."
    )
    _add_para(
        doc,
        "Data handling in the B2B context requires careful attention. When a school or hagwon "
        "provides student information to the platform for personalization, the school is the "
        "data controller and the platform is the data processor under PIPA. A formal data "
        "processing agreement (위탁계약) is required under PIPA Article 26, specifying the "
        "scope of processing, security measures, subcontracting restrictions, and liability "
        "allocation."
    )

    _risk_tag(doc, "Low-Medium")

    _sub_head(doc, "B.", "Content API Licensing Considerations")
    _add_para(
        doc,
        "The API content licensing revenue stream (licensing review data to streaming platforms "
        "and parental control apps at $10-50K/year per client) requires standard commercial "
        "licensing documentation. Key contractual considerations include:"
    )
    _add_bullet(doc, "Whether review data constitutes the platform's original intellectual property (it does, as transformative AI-generated analysis).")
    _add_bullet(doc, "Limitations on the licensee's right to modify, redistribute, or present the data out of context.")
    _add_bullet(doc, "Warranties regarding accuracy and reliability of AI-generated reviews and age ratings.")
    _add_bullet(doc, "Indemnification for claims arising from content errors or omissions.")
    _add_bullet(doc, "Data protection obligations where API data includes or relates to user data.")

    _risk_tag(doc, "Low")

    doc.add_page_break()


def section_other(doc):
    _section_head(doc, "VIII", "Other Risks")

    _sub_head(doc, "A.", "Sponsorship Conflict of Interest and Fair Trade Rules")
    _add_para(
        doc,
        "The platform's 'TripAdvisor model'—where studios fund the platform by paying for "
        "featured placement and engagement pack creation, while parents use the platform for "
        "free or at a subsidized price—creates an inherent tension between commercial incentives "
        "and review independence. Even if studios 'cannot pay to change review scores,' the "
        "commercial relationship creates an appearance of conflict that can undermine consumer "
        "trust and attract regulatory attention."
    )
    _add_para(
        doc,
        "Under Korean fair trade law, the KFTC may investigate whether the sponsorship "
        "arrangement constitutes an undisclosed commercial endorsement or whether it creates "
        "misleading impressions about the independence of content reviews. The KFTC's Digital "
        "Platform Competition Promotion Act (디지털 플랫폼 경쟁촉진법), expected to take effect "
        "in 2025-2026, may impose additional transparency obligations on digital platforms."
    )
    _risk_tag(doc, "Medium-High")

    _sub_head(doc, "B.", "Expansion Market Regulatory Overview")

    _sub_sub_head(doc, "(i)", "Japan")
    _add_para(
        doc,
        "Japan's Copyright Act Article 30-4 permits 'use of copyrighted works not for the "
        "purpose of personally enjoying or causing another person to enjoy the thoughts or "
        "sentiments expressed therein'—a provision specifically designed to enable AI and "
        "computational analysis. This provision is more permissive than either the U.S. or "
        "Korean fair-use frameworks for AI analysis purposes. Japan's Act on Personal "
        "Information Protection (APPI) requires consent for personal data collection from "
        "minors, though the age threshold is less prescriptive than COPPA or PIPA. The Act "
        "on Development of an Environment that Provides Safe and Secure Internet Use for "
        "Young People imposes filtering and safety obligations on internet service providers."
    )

    _sub_sub_head(doc, "(ii)", "Vietnam")
    _add_para(
        doc,
        "Vietnam's Cybersecurity Law (2018) and Decree 53/2022 impose mandatory data "
        "localization: enterprises providing services on telecommunication networks in Vietnam "
        "must store personal data locally with a minimum 24-month retention period and establish "
        "a physical presence (branch or representative office) in Vietnam. The Law on Children "
        "(2016) defines a child as under 16 and requires dual consent (both child age 7+ and "
        "parent/guardian) for data publication. Children's data is classified as sensitive under "
        "Decree 13/2023 (personal data protection decree), requiring enhanced protection and "
        "impact assessment for cross-border transfers."
    )

    _sub_sub_head(doc, "(iii)", "Indonesia")
    _add_para(
        doc,
        "Indonesia's Personal Data Protection Law (UU PDP, Law No. 27 of 2022) classifies "
        "children's data as 'specific (sensitive) personal data' requiring parental or guardian "
        "consent. However, the PDP Law does not define the age threshold for 'child,' creating "
        "significant uncertainty—the Child Protection Law (Law No. 35 of 2014) defines a child "
        "as under 18, but no authoritative guidance confirms this cross-reference for PDP "
        "purposes. [Unverified] Critically, despite the two-year transitional period expiring in "
        "October 2024, the implementing regulations remain in draft as of March 2026. "
        "Administrative penalties can reach 2% of annual revenue. The supervisory body has not "
        "yet been formed. This regulatory uncertainty represents a material risk factor for "
        "market entry timing."
    )

    _risk_tag(doc, "Medium")
    _add_para(doc, "Recommended Mitigation:", bold=True)
    _add_bullet(doc, "Commission jurisdiction-specific legal opinions before entering each expansion market.")
    _add_bullet(doc, "Japan's copyright framework is favorable; prioritize Japan as the first expansion market.")
    _add_bullet(doc, "For Vietnam and Indonesia, plan for data localization and local representative requirements.")

    _sub_head(doc, "C.", "Government Grant Conditions (TIPS Program)")
    _add_para(
        doc,
        "The TIPS (Tech Incubator Program for Startups) program, administered by the Ministry "
        "of SMEs and Startups, provides up to KRW 500 million (approximately USD 370,000) in "
        "R&D funding plus commercialization support. Legal conditions typically include:"
    )
    _add_bullet(doc, "Funded IP is generally owned by the startup, but the government retains a non-exclusive license for government purposes and may restrict transfer of core technology to foreign entities.", bold_prefix="IP ownership: ")
    _add_bullet(doc, "Technical and commercial milestones must be achieved within specified timelines; failure may trigger clawback provisions.", bold_prefix="Milestone obligations: ")
    _add_bullet(doc, "Grant funds must be used for specified R&D purposes; diversion to general operating expenses or non-approved activities constitutes misuse.", bold_prefix="Use restrictions: ")
    _add_bullet(doc, "Periodic financial and technical progress reports are mandatory. Audit rights are standard.", bold_prefix="Reporting: ")
    _add_bullet(doc, "Some government programs restrict foreign investment exceeding certain thresholds during the funding period or require prior approval for technology transfer abroad.", bold_prefix="Foreign expansion: ")

    _risk_tag(doc, "Low-Medium")

    doc.add_page_break()


def risk_matrix(doc):
    _section_head(doc, "IX", "Consolidated Risk Summary Matrix")

    _add_table(
        doc,
        ["#", "Risk Area", "Risk Level", "Likelihood", "Impact", "Priority"],
        [
            ["1", "Screenshots / frame captures", "High", "High", "High", "Immediate"],
            ["2", "OTT platform ToS violations", "High", "High", "High", "Immediate"],
            ["3", "Children's data (KR PIPA under-14)", "High", "High", "High", "Pre-launch"],
            ["4", "Sponsorship disclosure / conflict", "Medium-High", "Medium", "High", "Pre-launch"],
            ["5", "Children's data (US COPPA)", "Medium-High", "Medium", "High", "Pre-US launch"],
            ["6", "Children's data (EU GDPR/AADC)", "Medium-High", "Medium", "High", "Pre-EU launch"],
            ["7", "Copyright (AI analysis & summaries)", "Medium", "Medium", "Medium", "Ongoing"],
            ["8", "AI regulation (EU AI Act)", "Medium", "Medium", "Medium", "By Aug 2026"],
            ["9", "Gamification / nudge patterns", "Medium", "Medium", "Medium", "Design phase"],
            ["10", "Subtitle/script copyrightability", "Medium", "Low", "High", "Design phase"],
            ["11", "Expansion market regulations", "Medium", "Medium", "Medium", "Pre-expansion"],
            ["12", "B2B education compliance", "Low-Medium", "Low", "Medium", "Pre-B2B launch"],
            ["13", "Government grant (TIPS) conditions", "Low-Medium", "Low", "Low", "Ongoing"],
            ["14", "API content licensing", "Low", "Low", "Low", "Contract phase"],
        ],
    )


def counter_analysis(doc):
    _section_head(doc, "X", "Counter-Analysis and Alternative Interpretations")
    _add_para(
        doc,
        "In the interest of completeness, we note the following counter-arguments and "
        "alternative interpretations that qualify the conclusions above:"
    )

    _add_para(doc, "Copyright fair use may be weaker than presented.", bold=True)
    _add_para(
        doc,
        "If courts adopt a narrow reading of Andy Warhol Foundation v. Goldsmith, the "
        "commercial nature of the platform and its systematic processing of entire catalogs "
        "of copyrighted content could weigh against fair use, even for transformative outputs. "
        "The fact that the platform analyzes content at industrial scale (processing 'all episodes "
        "of all shows') distinguishes it from individual critical commentary."
    )

    _add_para(doc, "ToS enforcement risk may be lower than assessed.", bold=True)
    _add_para(
        doc,
        "In practice, major streaming platforms have limited incentive to pursue legal action "
        "against a platform that drives viewership and engagement. The 'content discovery' value "
        "proposition—where the platform helps parents find and engage with content—may align "
        "with platform interests. Several precedents in the travel industry (TripAdvisor, "
        "Yelp) demonstrate that review platforms can coexist with service providers despite "
        "initial friction."
    )

    _add_para(doc, "Children's data scope may be narrower than assumed.", bold=True)
    _add_para(
        doc,
        "If the platform is strictly designed so that parents create accounts, manage all "
        "data, and children never directly interact with the platform, the regulatory "
        "classification as a 'children's service' may be avoidable. However, this depends "
        "on precise implementation details—any child-facing feature (gamification badges, "
        "personalized recommendations based on child's age) could trigger children's data "
        "obligations."
    )

    _add_para(doc, "EU AI Act classification uncertainty.", bold=True)
    _add_para(
        doc,
        "The boundary between 'limited risk' and 'high risk' under the EU AI Act is not yet "
        "settled by regulatory guidance. If the European Commission or national authorities "
        "classify AI-based educational content recommendations for children as high-risk "
        "(under Annex III, Category 3), compliance costs and timelines would increase "
        "substantially."
    )

    doc.add_page_break()


def practical_implications(doc):
    _section_head(doc, "XI", "Practical Implications and Recommended Action Plan")
    _add_para(doc, "Based on the foregoing analysis, we recommend the following prioritized actions:", bold=True)

    _add_para(doc, "Phase 1 — Pre-Launch (Months 1-6):", bold=True)
    _add_bullet(doc, "Implement PIPA-compliant parental consent system with identity verification for under-14 data collection.")
    _add_bullet(doc, "Establish an 'AI vision only' content analysis pipeline that does not access subtitle files, APIs, or backend data.")
    _add_bullet(doc, "Prohibit screenshots and frame captures in all user-facing content; use text descriptions and original illustrations only.")
    _add_bullet(doc, "Draft and adopt an editorial independence charter separating content analysis from commercial partnerships.")
    _add_bullet(doc, "Prepare a children's privacy policy and cookie/tracking policy for the Korean market.")

    _add_para(doc, "Phase 2 — Korean Market Operations (Months 6-18):", bold=True)
    _add_bullet(doc, "Negotiate content analysis licensing agreements with at least one major Korean studio (e.g., ICONIX, CJ ENM).")
    _add_bullet(doc, "Implement KFTC-compliant endorsement disclosure for all studio-sponsored content.")
    _add_bullet(doc, "Designate a youth protection officer (청소년 보호 책임자) if daily user thresholds are reached.")
    _add_bullet(doc, "Conduct ongoing bias audits of AI content analysis outputs.")

    _add_para(doc, "Phase 3 — International Expansion (Year 2+):", bold=True)
    _add_bullet(doc, "Commission jurisdiction-specific legal opinions for each expansion market (Japan, U.S., EU, Vietnam, Indonesia) at least 6 months before market entry.")
    _add_bullet(doc, "Implement COPPA compliance for U.S. market entry; conduct DPIA for EU market entry.")
    _add_bullet(doc, "Prepare for EU AI Act transparency obligations (compliance date: August 2, 2026).")
    _add_bullet(doc, "Establish data processing agreements for B2B school/hagwon clients.")

    doc.add_page_break()


def bibliography(doc):
    _section_head(doc, "XII", "Annotated Bibliography — Key Authorities")
    _add_para(doc, "Primary Sources — Legislation [P#]:", bold=True)

    sources = [
        ("[P1]", "Korean Copyright Act (저작권법), Art. 2, 28, 35-3, 35-5, 136 — effective Sept. 26, 2025. Korea's primary copyright statute including fair-use provision."),
        ("[P2]", "Korean Personal Information Protection Act (PIPA, 개인정보 보호법), Art. 22-2, 26, 71 — effective Oct. 2, 2025. Children's data consent requirements and penalties."),
        ("[P3]", "Korean Youth Protection Act (청소년 보호법), Art. 2, 7, 9 — effective Oct. 1, 2025. Youth harmful media regulation framework."),
        ("[P4]", "Korean Network Act (정보통신망법), Art. 42, 42-2, 42-3, 44 — effective Oct. 1, 2025. Online content and youth protection obligations."),
        ("[P5]", "Korean Fair Labeling and Advertising Act (표시·광고의 공정화에 관한 법률) — effective Jan. 21, 2025. Endorsement disclosure requirements."),
        ("[P6]", "U.S. Copyright Act, 17 U.S.C. § 107 (Fair Use)."),
        ("[P7]", "U.S. Children's Online Privacy Protection Act (COPPA), 15 U.S.C. §§ 6501–6506."),
        ("[P8]", "EU General Data Protection Regulation (GDPR), Regulation (EU) 2016/679, Art. 6, 8, Recitals 38, 58."),
        ("[P9]", "EU Copyright Directive, Directive (EU) 2019/790, Art. 3, 4."),
        ("[P10]", "EU AI Act, Regulation (EU) 2024/1689, Art. 6, 50, Annex III."),
        ("[P11]", "EU Digital Services Act, Regulation (EU) 2022/2065, Art. 25 (dark patterns)."),
        ("[P12]", "Japanese Copyright Act, Art. 30-4 (computational analysis exception)."),
    ]
    for tag, desc in sources:
        _add_bullet(doc, f" {desc}", bold_prefix=tag)

    _add_para(doc, "Case Law [C#]:", bold=True)
    cases = [
        ("[C1]", "Authors Guild v. Google, Inc., 804 F.3d 202 (2d Cir. 2015) — Google Books digitization held 'highly transformative' fair use."),
        ("[C2]", "Andy Warhol Foundation for the Visual Arts v. Goldsmith, 598 U.S. 508 (2023) — Narrowed transformative-use doctrine; commercial use with same purpose weighs against fair use."),
        ("[C3]", "Thomson Reuters v. Ross Intelligence Inc., No. 1:20-cv-613 (D. Del. Feb. 11, 2025) — First court to reject fair use for AI training on copyrighted text; market substitute analysis decisive."),
        ("[C4]", "Bartz v. Anthropic PBC, No. 3:23-cv-07721 (N.D. Cal. June 23, 2025) — AI training on copyrighted books held 'quintessentially transformative' fair use."),
        ("[C5]", "Warner Bros. v. RDR Books, 575 F. Supp. 2d 513 (S.D.N.Y. 2008) — Reference guide held infringing for reproducing too much original expression."),
        ("[C6]", "Kelly v. Arriba Soft Corp., 336 F.3d 811 (9th Cir. 2003) — Thumbnail reproductions in image search held fair use."),
        ("[C7]", "Van Buren v. United States, 593 U.S. 374 (2021) — CFAA requires access without authorization, not merely exceeding authorized use."),
        ("[C8]", "MDY Industries v. Blizzard Entertainment, Inc. (9th Cir. 2010) — ToS violation is not copyright infringement absent nexus to exclusive rights."),
        ("[C9]", "Reddit v. Perplexity AI (2025); NYT v. Perplexity AI (Dec. 2025) — DMCA Section 1201 anti-circumvention theories against AI services."),
    ]
    for tag, desc in cases:
        _add_bullet(doc, f" {desc}", bold_prefix=tag)

    _add_para(doc, "Administrative / Regulatory Guidance [A#]:", bold=True)
    admin = [
        ("[A1]", "Korean PIPC — Guidelines on Collection and Use of Children's Personal Information."),
        ("[A2]", "Korean KFTC — Guidelines on Endorsement and Recommendation Advertising (추천·보증 표시·광고 심사지침)."),
        ("[A3]", "U.S. FTC — 16 C.F.R. Part 312 (COPPA Rule), as amended 2024."),
        ("[A4]", "U.S. FTC — 16 C.F.R. Part 255 (Endorsement Guides), updated June 2023."),
        ("[A5]", "UK ICO — Age Appropriate Design Code (Children's Code), 15 standards, effective September 2021."),
    ]
    for tag, desc in admin:
        _add_bullet(doc, f" {desc}", bold_prefix=tag)


def verification_guide(doc):
    _section_head(doc, "XIII", "Verification Guide")
    _add_para(
        doc,
        "The following resources may be consulted to verify the legal authorities cited in "
        "this memorandum:"
    )
    _add_table(
        doc,
        ["Source Type", "Verification Resource", "URL / Access Method"],
        [
            ["Korean statutes", "National Law Information Center", "law.go.kr (Official — Grade A)"],
            ["Korean statutes (English)", "Korea Legislation Research Institute", "elaw.klri.re.kr (Unofficial translation — Grade B max)"],
            ["Korean case law", "Supreme Court Comprehensive Legal Information", "supremecourt.go.kr"],
            ["Korean regulatory guidance", "PIPC (개인정보보호위원회)", "pipc.go.kr"],
            ["Korean advertising rules", "KFTC (공정거래위원회)", "ftc.go.kr"],
            ["U.S. statutes / case law", "Congress.gov, Google Scholar, Westlaw", "congress.gov"],
            ["U.S. FTC guidance", "Federal Trade Commission", "ftc.gov"],
            ["EU legislation", "EUR-Lex", "eur-lex.europa.eu"],
            ["UK AADC", "ICO (Information Commissioner's Office)", "ico.org.uk/for-organisations/childrens-code-hub"],
            ["Platform ToS", "Netflix, Disney+, YouTube", "Direct review of current terms at each platform"],
        ],
    )
    _add_para(
        doc,
        "Note: All Korean statutes cited were retrieved via the Open Law API (law.go.kr DRF "
        "API) on March 26, 2026. Effective dates are as stated in each statute's enforcement "
        "provisions. U.S. and EU authorities are based on published legislation and case law "
        "available as of the same date.",
        color=COLOR_MUTED,
        size_pt=FONT_SIZE_SMALL,
    )

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(closing.add_run("— End of Memorandum —"), bold=True, color=COLOR_TITLE, size_pt=12)


# ── Main ─────────────────────────────────────────────────────────
def main():
    doc = Document()
    setup(doc)
    cover(doc)
    toc(doc)
    executive_summary(doc)
    section_copyright(doc)
    section_tos(doc)
    section_privacy(doc)
    section_ai_regulation(doc)
    section_child_services(doc)
    section_b2b(doc)
    section_other(doc)
    risk_matrix(doc)
    counter_analysis(doc)
    practical_implications(doc)
    bibliography(doc)
    verification_guide(doc)

    out = Path("output/CCEP_Legal_Risk_Review_2026-03-26.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Saved → {out}")


if __name__ == "__main__":
    main()
