from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FONT_NAME = "Batang"
FONT_SIZE_PT = 10
COLOR_TITLE = RGBColor(25, 55, 109)
COLOR_MUTED = RGBColor(96, 96, 96)
COLOR_ALERT = RGBColor(170, 20, 20)
TABLE_HEADER_FILL = "E7EEF8"


@dataclass
class Meta:
    confidentiality: str = ""
    client: str = ""
    matter: str = ""
    as_of_date: str = ""
    issue_date: str = ""
    language: str = ""


def set_run_font(run, *, bold: bool | None = None, size_pt: float | None = None, color=None):
    run.font.name = FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size_pt if size_pt is not None else FONT_SIZE_PT)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_field(paragraph, instruction: str):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    paragraph._p.append(fld)


def set_paragraph_border_bottom(paragraph, color_hex: str = "C0C0C0", size: str = "6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)


def shade_cell(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def apply_doc_style(doc: Document):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(FONT_SIZE_PT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(5)

    for name, size, before, after in [
        ("Heading 1", 10, 12, 6),
        ("Heading 2", 10, 10, 5),
        ("Heading 3", 10, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = COLOR_TITLE
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    list_bullet = doc.styles["List Bullet"]
    list_bullet.font.name = FONT_NAME
    list_bullet._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    list_bullet.font.size = Pt(FONT_SIZE_PT)

    list_number = doc.styles["List Number"]
    list_number.font.name = FONT_NAME
    list_number._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    list_number.font.size = Pt(FONT_SIZE_PT)


def setup_header_footer(section):
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("FORMAL LEGAL OPINION | CONFIDENTIAL")
    set_run_font(run, bold=True, size_pt=10, color=COLOR_MUTED)

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("Page ")
    set_run_font(r1, size_pt=10, color=COLOR_MUTED)
    add_field(fp, "PAGE")
    r2 = fp.add_run(" of ")
    set_run_font(r2, size_pt=10, color=COLOR_MUTED)
    add_field(fp, "NUMPAGES")


def tune_paragraph(paragraph, *, justify: bool = True, indent_cm: float = 0.0):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.first_line_indent = Cm(indent_cm)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run)


def add_inline_runs(paragraph, text: str):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def parse_meta(lines: list[str]) -> tuple[Meta, int]:
    meta = Meta()
    start = 0
    for i, line in enumerate(lines):
        if line.strip().startswith("# Formal Legal Opinion Letter"):
            start = i + 1
            break
    mapping = {
        "Confidentiality": "confidentiality",
        "Client": "client",
        "Matter": "matter",
        "As-of Date": "as_of_date",
        "Issue Date": "issue_date",
        "Language": "language",
    }
    idx = start
    while idx < len(lines):
        s = lines[idx].strip()
        if s.startswith("**Scope and Assumptions**"):
            break
        m = re.match(r"-\s+\*\*([^*]+)\*\*:\s*(.+)", s)
        if m:
            key = m.group(1).strip()
            val = m.group(2).strip()
            attr = mapping.get(key)
            if attr:
                setattr(meta, attr, val)
        idx += 1
    return meta, idx


def add_cover(doc: Document, meta: Meta):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(meta.confidentiality or "Attorney Work Product / Client Internal Use Only")
    set_run_font(run, bold=True, size_pt=10, color=COLOR_ALERT)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("FORMAL LEGAL OPINION LETTER")
    set_run_font(t_run, bold=True, size_pt=10, color=COLOR_TITLE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = subtitle.add_run("Brazil ECA Digital (Lei no. 15.211/2025) Impact Assessment")
    set_run_font(s_run, size_pt=10, color=COLOR_MUTED)
    set_paragraph_border_bottom(subtitle, color_hex="B7C6E6", size="8")

    doc.add_paragraph()
    info = doc.add_table(rows=5, cols=2)
    info.style = "Table Grid"
    info.autofit = True
    rows = [
        ("Client", meta.client),
        ("Matter", meta.matter),
        ("As-of Date", meta.as_of_date),
        ("Issue Date", meta.issue_date),
        ("Language", meta.language),
    ]
    for i, (k, v) in enumerate(rows):
        left = info.rows[i].cells[0]
        right = info.rows[i].cells[1]
        left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(left, TABLE_HEADER_FILL)
        lp = left.paragraphs[0]
        lr = lp.add_run(k)
        set_run_font(lr, bold=True)
        tune_paragraph(lp, justify=False)
        rp = right.paragraphs[0]
        rr = rp.add_run(v)
        set_run_font(rr)
        tune_paragraph(rp, justify=False)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    n_run = note.add_run(
        "Disclaimer: 본 문서는 공개 법령 및 공적 자료를 기초로 한 일반적 법률의견서로서, "
        "개별 분쟁 또는 사실관계에 대한 확정적 소송의견을 대체하지 않습니다."
    )
    set_run_font(n_run, size_pt=10, color=COLOR_MUTED)

    doc.add_page_break()


def add_toc(doc: Document):
    h = doc.add_heading("Table of Contents", level=1)
    for r in h.runs:
        set_run_font(r, bold=True, size_pt=10, color=COLOR_TITLE)
    p = doc.add_paragraph()
    add_field(p, r'TOC \\o "1-3" \\h \\z \\u')
    tune_paragraph(p, justify=False)
    hint = doc.add_paragraph(
        "Note: 목차 번호가 비어 보이면 Word에서 문서를 열어 'Update Field'를 실행하십시오."
    )
    tune_paragraph(hint, justify=False)
    for r in hint.runs:
        set_run_font(r, size_pt=10, color=COLOR_MUTED)
    doc.add_page_break()


def is_md_table_line(line: str) -> bool:
    t = line.strip()
    return t.startswith("|") and t.endswith("|") and "|" in t[1:-1]


def is_md_separator(line: str) -> bool:
    t = line.strip().strip("|").replace(" ", "")
    return bool(t) and all(c in "-:|" for c in t)


def split_md_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def render_md_table(doc: Document, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    table.autofit = True

    for c, text in enumerate(header):
        cell = table.rows[0].cells[c]
        shade_cell(cell, TABLE_HEADER_FILL)
        p = cell.paragraphs[0]
        p.text = ""
        r = p.add_run(text)
        set_run_font(r, bold=True)
        tune_paragraph(p, justify=False)

    for r_idx, row in enumerate(rows, start=1):
        row = row + [""] * (len(header) - len(row))
        for c, text in enumerate(row[: len(header)]):
            cell = table.rows[r_idx].cells[c]
            if r_idx % 2 == 0:
                shade_cell(cell, "F9F9F9")
            p = cell.paragraphs[0]
            p.text = ""
            add_inline_runs(p, text)
            tune_paragraph(p, justify=False)

    spacer = doc.add_paragraph()
    tune_paragraph(spacer, justify=False)


def render_body(doc: Document, lines: list[str], start_idx: int):
    i = start_idx
    while i < len(lines):
        raw = lines[i].rstrip()
        s = raw.strip()

        if not s:
            if i + 1 < len(lines) and lines[i + 1].strip() == "":
                i += 1
                continue
            p = doc.add_paragraph("")
            tune_paragraph(p, justify=False)
            i += 1
            continue

        if s.startswith("[포맷팅 지시:"):
            p = doc.add_paragraph()
            add_inline_runs(p, s)
            tune_paragraph(p, justify=False)
            for r in p.runs:
                set_run_font(r, size_pt=10, color=COLOR_MUTED)
            i += 1
            continue

        if is_md_table_line(s) and i + 1 < len(lines) and is_md_separator(lines[i + 1]):
            header = split_md_row(s)
            rows = []
            i += 2
            while i < len(lines) and is_md_table_line(lines[i].strip()):
                rows.append(split_md_row(lines[i]))
                i += 1
            render_md_table(doc, header, rows)
            continue

        if s.startswith("# "):
            p = doc.add_heading(s[2:].strip(), level=1)
            tune_paragraph(p, justify=False)
            set_paragraph_border_bottom(p, color_hex="B7C6E6", size="6")
            i += 1
            continue
        if s.startswith("## "):
            p = doc.add_heading(s[3:].strip(), level=2)
            tune_paragraph(p, justify=False)
            i += 1
            continue
        if s.startswith("### "):
            p = doc.add_heading(s[4:].strip(), level=3)
            tune_paragraph(p, justify=False)
            i += 1
            continue

        if re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, s)
            tune_paragraph(p, justify=True)
            i += 1
            continue

        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, s[2:].strip())
            tune_paragraph(p, justify=True)
            i += 1
            continue

        if re.fullmatch(r"\*\*[^*]+\*\*", s):
            p = doc.add_paragraph()
            add_inline_runs(p, s)
            tune_paragraph(p, justify=False)
            for r in p.runs:
                set_run_font(r, bold=True, color=COLOR_TITLE)
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, s)
        tune_paragraph(p, justify=True, indent_cm=0.0)
        i += 1


def build_professional_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    meta, body_start_idx = parse_meta(lines)

    doc = Document()
    apply_doc_style(doc)
    setup_header_footer(doc.sections[0])
    add_cover(doc, meta)
    add_toc(doc)
    render_body(doc, lines, body_start_idx)
    doc.save(docx_path)


def main():
    root = Path(__file__).resolve().parents[1]
    md_path = root / "output" / "reports" / "brazil_eca_digital_formal_legal_opinion_ko_2026-03-05.md"
    docx_path = root / "output" / "reports" / "brazil_eca_digital_formal_legal_opinion_ko_2026-03-05.docx"
    build_professional_docx(md_path, docx_path)
    print(docx_path)


if __name__ == "__main__":
    main()
