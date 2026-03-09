"""
Generate a professional law-firm-quality .docx comparative analysis:
Korean Attorney Act Article 26-2 (ACP) vs. US Attorney-Client Privilege.

Output: A4 .docx with Mode B (Comparative Matrix) + Mode D (Black-letter & Commentary).
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Design tokens ──────────────────────────────────────────────────────
FONT_BODY = "Times New Roman"
FONT_BODY_KO = "맑은 고딕"
FONT_HEADER = "Arial"
FONT_SIZE = Pt(10.5)
FONT_SIZE_SMALL = Pt(9)
FONT_SIZE_TINY = Pt(8)

COLOR_NAVY = RGBColor(0x19, 0x37, 0x6D)
COLOR_DARK = RGBColor(0x1B, 0x2A, 0x4A)
COLOR_BODY = RGBColor(0x00, 0x00, 0x00)
COLOR_MUTED = RGBColor(0x60, 0x60, 0x60)
COLOR_ALERT = RGBColor(0xAA, 0x14, 0x14)
COLOR_ACCENT = RGBColor(0x8B, 0x73, 0x55)

FILL_HEADER = "1B2A4A"
FILL_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
FILL_LIGHT = "EDF2F9"
FILL_ALT = "F7F9FC"
FILL_UNVERIFIED = "FFF8E1"


# ── Low-level helpers ──────────────────────────────────────────────────
def _set_cell_shading(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_cell_width(cell, cm_val: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(cm_val * 567)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def _set_run_font(run, *, font=None, size=None, bold=None, italic=None, color=None):
    f = font or FONT_BODY
    run.font.name = f
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:ascii"), f)
    rf.set(qn("w:hAnsi"), f)
    rf.set(qn("w:eastAsia"), FONT_BODY_KO)
    run.font.size = size or FONT_SIZE
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_border_bottom(paragraph, color_hex="1B2A4A", sz="6"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)


def _add_page_field(paragraph, field_code: str):
    """Insert a Word field (PAGE, NUMPAGES, TOC, etc.) using fldSimple."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), field_code)
    # Add an empty run inside so Word can populate it
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = ""
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def _add_inline(paragraph, text: str, *, base_color=COLOR_BODY):
    """Parse **bold**, *italic*, `code`, and [Unverified] markers."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[Unverified\]|\[Unresolved Conflict\])", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            _set_run_font(r, bold=True, color=base_color)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            r = paragraph.add_run(part[1:-1])
            _set_run_font(r, italic=True, color=base_color)
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            _set_run_font(r, font="Consolas", size=Pt(9.5), color=COLOR_MUTED)
        elif part in ("[Unverified]", "[Unresolved Conflict]"):
            r = paragraph.add_run(f" {part} ")
            _set_run_font(r, bold=True, size=Pt(9), color=COLOR_ALERT)
        else:
            r = paragraph.add_run(part)
            _set_run_font(r, color=base_color)


def _body_paragraph(doc, text, *, indent_cm=0.0, justify=True, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    _add_inline(p, text)
    return p


def _bullet(doc, text, *, indent_cm=1.0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("•  ")
    _set_run_font(r, color=COLOR_NAVY, bold=True)
    _add_inline(p, text)
    return p


def _numbered(doc, text, num: int, *, indent_cm=1.0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{num}.  ")
    _set_run_font(r, bold=True, color=COLOR_NAVY)
    _add_inline(p, text)
    return p


# ── Section heading helpers ────────────────────────────────────────────
def _heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    _set_run_font(r, font=FONT_BODY, size=Pt(13), bold=True, color=COLOR_NAVY)
    _add_border_bottom(p, color_hex="1B2A4A", sz="8")
    return p


def _heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    _set_run_font(r, font=FONT_BODY, size=Pt(11.5), bold=True, color=COLOR_DARK)
    return p


def _heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    _set_run_font(r, font=FONT_BODY, size=Pt(10.5), bold=True, color=RGBColor(0x33, 0x33, 0x33))
    return p


# ── Table helpers ──────────────────────────────────────────────────────
def _add_table(doc, headers, rows, *, col_widths_cm=None):
    """Create a styled table with dark header row."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for c, h_text in enumerate(headers):
        cell = table.rows[0].cells[c]
        _set_cell_shading(cell, FILL_HEADER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h_text)
        _set_run_font(r, font=FONT_BODY, size=Pt(9.5), bold=True, color=FILL_HEADER_TEXT)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        fill = FILL_ALT if row_idx % 2 == 1 else None
        for c in range(n_cols):
            cell = table.rows[row_idx + 1].cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if fill:
                _set_cell_shading(cell, fill)
            text = row_data[c] if c < len(row_data) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.2
            _add_inline(p, text, base_color=COLOR_BODY)

    # Column widths
    if col_widths_cm:
        for row in table.rows:
            for c, w in enumerate(col_widths_cm):
                if c < len(row.cells):
                    _set_cell_width(row.cells[c], w)

    # Spacer after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(2)
    return table


# ── Page setup ─────────────────────────────────────────────────────────
def _setup_page(doc):
    """A4, professional margins, default fonts."""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = FONT_SIZE
    style.font.color.rgb = COLOR_BODY
    rPr = style.element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.append(rf)
    rf.set(qn("w:ascii"), FONT_BODY)
    rf.set(qn("w:hAnsi"), FONT_BODY)
    rf.set(qn("w:eastAsia"), FONT_BODY_KO)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(6)


def _setup_header_footer(doc, header_text: str):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # First-page header: empty (cover page)
    # Subsequent pages
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(header_text)
    _set_run_font(r, font=FONT_HEADER, size=FONT_SIZE_TINY, bold=False, color=COLOR_MUTED)
    _add_border_bottom(p, color_hex="CCCCCC", sz="4")

    # Footer with page numbers
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("- ")
    _set_run_font(r1, font=FONT_HEADER, size=FONT_SIZE_TINY, color=COLOR_MUTED)
    _add_page_field(fp, "PAGE")
    r2 = fp.add_run(" -")
    _set_run_font(r2, font=FONT_HEADER, size=FONT_SIZE_TINY, color=COLOR_MUTED)


# ── Cover page ─────────────────────────────────────────────────────────
def _add_cover(doc):
    # Firm name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("법무법인 진주")
    _set_run_font(r, font=FONT_HEADER, size=Pt(22), bold=True, color=COLOR_NAVY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run("JINJU LAW FIRM")
    _set_run_font(r2, font=FONT_HEADER, size=Pt(11), bold=False, color=COLOR_MUTED)
    _add_border_bottom(p2, color_hex="8B7355", sz="10")

    # Spacer
    for _ in range(3):
        doc.add_paragraph()

    # Document type
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run("비교법 분석 보고서")
    _set_run_font(rt, font=FONT_BODY, size=Pt(14), bold=True, color=COLOR_NAVY)

    pt2 = doc.add_paragraph()
    pt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt2 = pt2.add_run("COMPARATIVE LEGAL ANALYSIS")
    _set_run_font(rt2, font=FONT_HEADER, size=Pt(10), bold=False, color=COLOR_MUTED)

    doc.add_paragraph()

    # Title
    pt3 = doc.add_paragraph()
    pt3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pt3.paragraph_format.line_spacing = 1.4
    rt3 = pt3.add_run(
        "개정 변호사법 제26조의2\n"
        "(변호사-의뢰인 비밀유지특권)\n"
        "vs.\n"
        "미국 Attorney-Client Privilege"
    )
    _set_run_font(rt3, font=FONT_BODY, size=Pt(13), bold=True, color=COLOR_DARK)

    doc.add_paragraph()

    # Metadata table
    meta_items = [
        ("관할 (Jurisdictions)", "대한민국 / 미국 (연방)"),
        ("기준일 (As-of Date)", "2026-03-09"),
        ("한국법 상태", "국회 의결 (2026.1.29.) / 미시행 (공포 후 1년)"),
        ("미국법 상태", "현행법 (Common Law + FRE)"),
        ("출력 모드", "B (Comparative Matrix) + D (Black-letter & Commentary)"),
        ("작성자", "김재식 변호사 (5년차 Associate)"),
    ]
    tbl = doc.add_table(rows=len(meta_items), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta_items):
        lc = tbl.rows[i].cells[0]
        rc = tbl.rows[i].cells[1]
        _set_cell_shading(lc, FILL_LIGHT)
        _set_cell_width(lc, 5.0)
        _set_cell_width(rc, 10.0)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = lc.paragraphs[0]
        lr = lp.add_run(k)
        _set_run_font(lr, size=Pt(9.5), bold=True, color=COLOR_DARK)
        lp.paragraph_format.space_before = Pt(3)
        lp.paragraph_format.space_after = Pt(3)
        rp = rc.paragraphs[0]
        rr = rp.add_run(v)
        _set_run_font(rr, size=Pt(9.5), color=COLOR_BODY)
        rp.paragraph_format.space_before = Pt(3)
        rp.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # Disclaimer
    pd = doc.add_paragraph()
    pd.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rd = pd.add_run(
        "Disclaimer: 본 문서는 공개 법령 및 공적 자료를 기초로 한 법률 조사 보고서로서, "
        "개별 사안에 대한 법률자문을 구성하지 않습니다. 구체적 법적 판단은 해당 관할의 "
        "자격 있는 변호사와 상의하시기 바랍니다."
    )
    _set_run_font(rd, size=FONT_SIZE_SMALL, italic=True, color=COLOR_MUTED)

    doc.add_page_break()


# ── TOC ────────────────────────────────────────────────────────────────
def _add_toc(doc):
    _heading1(doc, "목차 (Table of Contents)")

    toc_items = [
        ("I.", "분석 범위 및 전제 (Scope & Assumptions)"),
        ("II.", "비교 요약표 (Comparative Matrix)"),
        ("III.", "비교 축별 상세 분석 (Black-letter & Commentary)"),
        ("", "  1. 특권의 법적 근거 및 법원(法源)"),
        ("", "  2. 적용 주체(Privilege Holder) 및 인적 범위"),
        ("", "  3. 보호 대상 커뮤니케이션의 범위 및 요건"),
        ("", "  4. 특권의 예외·포기(Waiver) 사유"),
        ("", "  5. 제3자 공유 시 특권 유지 여부"),
        ("", "  6. 형사절차·수사기관 대응에서의 특권 행사"),
        ("", "  7. 위반 시 제재 및 구제수단"),
        ("", "  8. 시행 전 과도기 쟁점"),
        ("", "  9. 실무상 주요 쟁점 — 해석 공백 및 Gap 분석"),
        ("IV.", "용어 대조표 (Glossary)"),
        ("V.", "주석부 참고문헌 (Annotated Bibliography)"),
        ("VI.", "검증 가이드 (Verification Guide)"),
        ("VII.", "향후 확인 필요 사항 ([Unverified] 항목 종합)"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.line_spacing = 1.2
        if num:
            r1 = p.add_run(f"{num}  ")
            _set_run_font(r1, bold=True, size=Pt(10), color=COLOR_NAVY)
            r2 = p.add_run(title)
            _set_run_font(r2, size=Pt(10), color=COLOR_BODY)
        else:
            r2 = p.add_run(title)
            _set_run_font(r2, size=Pt(9.5), color=COLOR_MUTED)

    doc.add_page_break()


# ── MAIN CONTENT ───────────────────────────────────────────────────────
def _add_scope(doc):
    _heading1(doc, "I. 분석 범위 및 전제 (Scope & Assumptions)")

    _body_paragraph(doc, "**관할**: 대한민국 (국가 단위) / 미국 (연방 common law 중심, 주법 보충)")
    _body_paragraph(doc, "**기준일**: 2026-03-09")
    _body_paragraph(doc, "**한국법 상태**: 국회 본회의 의결 (2026. 1. 29.) / 공포 대기 또는 완료 / 미시행 (공포 후 1년 경과 시 시행, 2027년 상반기 예정)")
    _body_paragraph(doc, "**미국법 상태**: 현행법 (수백 년간 축적된 common law + Federal Rules of Evidence)")
    _body_paragraph(doc, "**구조적 비대칭**: 본 분석은 시행 전 제정법(한국)과 현행 판례법(미국)의 비교라는 본질적 비대칭 구조를 가지고 있습니다. 한국법 부분은 판례·유권해석이 부재하므로 조문 문언 분석 및 입법자료 기반 해석에 집중하였으며, 하위법령 미제정 사항은 [Unverified] 태그로 표시하였습니다.")

    _heading2(doc, "전제 및 한계")
    _numbered(doc, "한국 개정 변호사법 제26조의2는 2026. 1. 29. 국회 본회의를 통과하였으며, 공포 후 1년 경과 시 시행됩니다.", 1)
    _numbered(doc, "미국법은 연방 common law (Upjohn 기준) 및 Federal Rules of Evidence를 중심으로 분석하되, 주법 차이는 보충적으로 언급하였습니다.", 2)
    _numbered(doc, "개정 전 한국법에는 Attorney-Client Privilege에 해당하는 명시적 제정법 조항이 없었습니다. 기존 변호사법 제26조(비밀유지의무)와의 관계를 배경으로 서술하였습니다.", 3)
    _numbered(doc, "하위법령(대통령령·대법원규칙) 미제정 사항은 [Unverified] 태그로 처리하고, 향후 확인 필요 사항으로 별도 정리하였습니다.", 4)


def _add_comparative_matrix(doc):
    _heading1(doc, "II. 비교 요약표 (Comparative Matrix)")

    _body_paragraph(doc, "아래 표는 9개 비교 축에 대한 한·미 양국법의 핵심 차이를 요약한 것입니다. 상세 분석은 제III장 이하를 참조하여 주시기 바랍니다.")

    headers = ["비교 축", "한국 — 개정 변호사법 제26조의2 (미시행)", "미국 — Attorney-Client Privilege (현행)"]

    rows = [
        [
            "1. 법적 근거 및 법원(法源)",
            "제정법 (변호사법 제26조의2 신설) — 명문 입법으로 특권 창설",
            "Common law (판례법) + FRE Rule 501; 각 주 증거법/판례; Restatement (Third)",
        ],
        [
            "2. 적용 주체 및 인적 범위",
            "'변호사와 의뢰인 또는 의뢰인이 되려는 자' — 사내변호사·외국변호사 명시 규정 없음 [Unverified]",
            "Client(자연인·법인), Attorney(사내변호사 포함 — Upjohn), 보조인력(paralegal, 전문가 등)",
        ],
        [
            "3. 보호 대상 커뮤니케이션",
            "①항: 법률조력 목적 비밀 의사교환; ②항: 소송·수사·조사 위한 서류·자료(전자적 형태 포함)",
            "법률자문 목적 비밀 커뮤니케이션 (4요소: communication, privileged persons, confidence, legal purpose); Work Product Doctrine (별도 법리)",
        ],
        [
            "4. 예외·포기(Waiver)",
            "③항: 의뢰인 승낙, 공범관계·위법행위 관여, 중대한 공익상 필요, 변호사-의뢰인 분쟁, 법률 특별규정",
            "Voluntary waiver, Implied/At-issue waiver, Subject matter waiver, Inadvertent disclosure (FRE 502), Crime-fraud exception",
        ],
        [
            "5. 제3자 공유 시 특권 유지",
            "명문 규정 없음 [Unverified]",
            "Common Interest Doctrine, Joint Defense Privilege — 공동 법적 이해관계 시 특권 유지",
        ],
        [
            "6. 형사절차에서의 행사",
            "수사기관 압수수색에서 보호; 구체적 행사 절차 미확정 [Unverified]",
            "Privilege log, In camera review, Filter/Taint team, Fed. R. Crim. P. 41(g), Protective order",
        ],
        [
            "7. 위반 시 제재·구제",
            "독자적 제재 규정 미포함 [Unverified]; 기존 형법 §317, 변호사법 §26 등 별도 적용",
            "증거배제, 변호사 징계, 민사 손해배상, Disqualification, Contempt of court",
        ],
        [
            "8. 과도기 쟁점",
            "소급 적용 명시; 기존 압수수색 실무 충돌; 하위법령 제정 동향 미확정 [Unverified]",
            "N/A (현행법, 안정적 운영)",
        ],
        [
            "9. 실무 쟁점·Gap",
            "사내변호사 적용 불명확, 예외 조항 광범성, 국제소송 상호인정, 전자적 커뮤니케이션 선별",
            "성숙한 판례법; dual-purpose communication (In re Kellogg Brown & Root) 등",
        ],
    ]

    _add_table(doc, headers, rows, col_widths_cm=[2.8, 6.0, 6.0])


def _add_detailed_analysis(doc):
    _heading1(doc, "III. 비교 축별 상세 분석 (Black-letter & Commentary)")

    # ── Axis 1 ──
    _heading2(doc, "1. 특권의 법적 근거 및 법원(法源)")

    _heading3(doc, "한국 — 개정 변호사법 제26조의2")
    _body_paragraph(doc,
        "개정 전 한국법에는 Attorney-Client Privilege에 해당하는 명시적 제정법 조항이 존재하지 않았습니다. "
        "기존 변호사법 제26조는 \"변호사 또는 변호사이었던 자는 그 직무상 알게 된 비밀을 누설하여서는 아니 된다\"고 "
        "규정하여, 변호사에게 **비밀유지'의무'(duty of confidentiality)**를 부과하였으나, 이는 변호사의 윤리적·직업적 "
        "의무일 뿐 의뢰인이 주장할 수 있는 증거법상 **'특권'(privilege)**은 아니었습니다. [P2]"
    )
    _body_paragraph(doc,
        "2026. 1. 29. 국회 본회의를 통과한 개정 변호사법은 제26조의2를 신설하여, 변호사-의뢰인 비밀유지**권**(權)을 "
        "처음으로 명문화하였습니다. 이로써 한국은 의무(obligation) 기반 체계에서 권리(right/privilege) 기반 체계로 "
        "전환하는 전환점을 맞이하게 되었습니다. [P1, A1]"
    )

    _heading3(doc, "기존 제26조와 신설 제26조의2의 관계")
    _bullet(doc, "**제26조 (비밀유지의무)**: 변호사의 의무 — 직무상 비밀 누설 금지 (위반 시 형사처벌 가능: 형법 제317조 업무상비밀누설죄)")
    _bullet(doc, "**제26조의2 (비밀유지특권)**: 변호사와 의뢰인 쌍방의 권리 — 비밀인 의사교환 내용 및 서류의 공개를 거부할 수 있는 적극적 권리")
    _bullet(doc, "양 조항은 병존하며, 제26조의2는 제26조의 보호를 증거법적 차원으로 확장·보완하는 것으로 해석됩니다. [S1]")

    _heading3(doc, "미국 — Common Law + Federal Rules of Evidence")
    _body_paragraph(doc,
        "미국의 Attorney-Client Privilege는 \"the oldest of the privileges for confidential communications known "
        "to the common law\"로서, 16세기 영국법에서 기원하여 수백 년간 판례를 통해 발전하여 왔습니다. [C1]"
    )
    _body_paragraph(doc,
        "연방 차원에서는 Federal Rules of Evidence Rule 501이 근거 규범입니다: \"The common law — as interpreted by "
        "United States courts in the light of reason and experience — governs a claim of privilege.\" 즉, 연방법원에서는 "
        "common law 원칙이 직접 적용되며, 개별 특권을 열거하는 방식이 아닌 판례법적 발전에 위임하는 구조입니다. [P3]"
    )
    _body_paragraph(doc,
        "FRE Rule 502는 2008년 제정되어 Attorney-Client Privilege 및 Work Product의 **포기(waiver)에 대한 제한**을 규정하고 있습니다. "
        "주(州) 차원에서는 각 주의 증거법이 적용되며, California Evidence Code §950-962 등 제정법으로 특권을 규정하는 주와 "
        "common law에 의존하는 주가 혼재하고 있습니다. [P3, P4, S5]"
    )

    _heading3(doc, "핵심 차이")
    _body_paragraph(doc,
        "한국은 단일 제정법 조항으로 특권을 창설한 반면, 미국은 common law 판례의 축적 위에 제한적 제정법(FRE 501, 502)이 "
        "보충하는 다층적 구조를 갖추고 있습니다. 한국의 제정법 접근은 법적 안정성과 예측가능성 측면에서 장점을 가지나, 구체적 사안에 대한 "
        "유연한 판례법적 발전이 아직 이루어지지 않았다는 한계가 있습니다."
    )

    # ── Axis 2 ──
    _heading2(doc, "2. 적용 주체(Privilege Holder) 및 인적 범위")

    _heading3(doc, "한국")
    _body_paragraph(doc,
        "제26조의2 제1항은 \"변호사와 의뢰인 또는 의뢰인이 되려는 자(이하 '의뢰인등')\"를 적용 주체로 규정하고 있습니다. [P1]"
    )
    _bullet(doc, "**'변호사'의 범위**: 변호사법상 '변호사'는 대한변호사협회에 등록한 자를 의미합니다(변호사법 제7조). 사내변호사(企業内辯護士, in-house counsel)가 변호사 등록을 유지하는 한 형식적으로는 포함되나, 사내변호사의 법률자문과 경영판단이 혼합된 커뮤니케이션에 대한 특권 적용 기준은 조문상 불명확합니다. [Unverified]")
    _bullet(doc, "**외국변호사(Foreign Legal Consultant)**: 외국법자문사법에 따라 등록한 외국법자문사의 자문에 대한 적용 여부는 명문 규정이 없습니다. [Unverified]")
    _bullet(doc, "**의뢰인이 되려는 자(Prospective Client)**: 명시적으로 포함되어 있으며, 미국법의 prospective client 보호와 유사한 구조입니다.")

    _heading3(doc, "미국")
    _bullet(doc, "**Client**: 자연인 및 법인 모두 포함됩니다. 법인의 경우 *Upjohn Co. v. United States*, 449 U.S. 383 (1981)이 핵심 판례로, 종전의 'control group test'를 배제하고, 법률자문 목적으로 경영진 지시에 따라 직무 범위 내에서 이루어진 모든 직원의 커뮤니케이션에 특권을 인정하였습니다. [C1]")
    _bullet(doc, "**Attorney**: 변호사 자격자 + 합리적으로 변호사라고 믿은 자에 대한 상담도 보호됩니다. **사내변호사** 명시 포함 — *Kellogg Brown & Root, Inc. v. United States*, 756 F.3d 754 (D.C. Cir. 2014)는 'in-house counsel status does not dilute the privilege'를 확인하였습니다. [S5]")
    _bullet(doc, "**Agents/Representatives**: 변호사가 법률자문을 위해 고용한 보조인력(paralegal, 회계사, 전문가 등)과의 커뮤니케이션도 보호 대상입니다. [S5, S6]")
    _bullet(doc, "**Prospective Client**: 변호사 선임에 이르지 않은 초기 상담도 특권 대상으로 보호됩니다. [S5]")

    _heading3(doc, "핵심 Gap")
    _body_paragraph(doc,
        "한국 개정법은 사내변호사 및 외국변호사의 적용 여부에 대한 명시적 규정이 없어, 향후 하위법령 또는 판례를 통한 해석이 "
        "필요합니다. 미국은 *Upjohn*과 후속 판례를 통해 사내변호사·법인 직원·보조인력까지 인적 범위를 체계적으로 확립하고 있습니다. [Unverified]"
    )

    # ── Axis 3 ──
    _heading2(doc, "3. 보호 대상 커뮤니케이션의 범위 및 요건")

    _heading3(doc, "한국 — 이원적 보호 구조")
    _body_paragraph(doc, "**① 의사교환 내용 (제1항 — ACP 본체)**")
    _body_paragraph(doc,
        "\"법률사건(法律事件) 또는 법률사무(法律事務)에 관한 조력을 제공하거나 받을 목적으로 이루어진 비밀인 의사교환 내용\"을 보호합니다. "
        "요건은 (i) 변호사-의뢰인등 간 의사교환, (ii) 법률 조력 목적, (iii) 비밀성이며, 비법률적 커뮤니케이션(경영 자문, 사업 판단 등)은 "
        "배제되는 것으로 해석됩니다. [S1]",
        indent_cm=0.5
    )
    _body_paragraph(doc, "**② 서류·자료 (제2항 — Work Product에 유사)**")
    _body_paragraph(doc,
        "\"수임한 사건과 관련하여 소송, 수사 또는 조사를 위하여 작성한 서류나 자료(전자적 형태로 작성, 관리되는 것을 포함)\"를 보호합니다. "
        "이는 미국법의 Work Product Doctrine과 유사한 기능을 하나, 별도의 독립 법리가 아닌 동일 조항 내에 규정되어 있습니다. [S1, S2]",
        indent_cm=0.5
    )

    _heading3(doc, "미국 — ACP + Work Product Doctrine (이원 법리)")
    _body_paragraph(doc, "**ACP의 4요소 (Wigmore 기준)**:")
    _numbered(doc, "**Communication**: 의뢰인과 변호사 간 의사교환 (구두·서면·전자적 형태 모두)", 1)
    _numbered(doc, "**Between privileged persons**: Counsel과 Client (또는 그 대리인) 간", 2)
    _numbered(doc, "**In confidence**: 비밀로 이루어질 것 — 제3자 동석 시 원칙적으로 상실", 3)
    _numbered(doc, "**For purpose of legal assistance**: 법률자문 또는 법률서비스 획득 목적", 4)

    _body_paragraph(doc,
        "**Work Product Doctrine** (별도 법리): *Hickman v. Taylor*, 329 U.S. 495 (1947)에서 확립되었습니다. "
        "소송 예상(anticipation of litigation) 하에 변호사가 작성한 자료를 보호하며, "
        "Ordinary work product (사실 자료) vs. Opinion work product (법적 판단·전략) — 후자가 더 강한 보호를 받습니다. [S5]"
    )

    _heading3(doc, "핵심 차이")
    _body_paragraph(doc,
        "한국은 ACP와 Work Product 유사 보호를 단일 조항(제1항·제2항)에 통합하였으나, 미국은 이를 별개의 독립 법리로 운영하며 "
        "각각 다른 요건·예외·보호 강도를 적용하고 있습니다. 특히 미국의 Work Product Doctrine은 상대방이 'substantial need'와 'undue "
        "hardship'을 입증하면 ordinary work product에 대한 개시를 명할 수 있으나, 한국 개정법 제2항에는 이러한 단계적 개시 "
        "메커니즘이 없습니다. [Unverified]"
    )

    # ── Axis 4 ──
    _heading2(doc, "4. 특권의 예외·포기(Waiver) 사유")

    _heading3(doc, "한국 — 제26조의2 제3항")
    _numbered(doc, "**의뢰인등의 승낙** (voluntary waiver): 의뢰인이 공개에 동의한 경우", 1)
    _numbered(doc, "**공범관계·위법행위 관여**: 변호사가 의뢰인등과 공범관계에 있거나, 증거인멸(證據湮滅), 범인은닉(犯人隱匿), 장물취득(贓物取得) 등 범죄·위법행위에 관여한 경우", 2)
    _numbered(doc, "**의뢰인의 위법 사용**: 의뢰인등이 의사교환 내용 또는 서류·자료를 위법행위에 사용하거나 사용하려고 한 경우 (미국의 crime-fraud exception과 유사)", 3)
    _numbered(doc, "**중대한 공익상 필요**: 위 2, 3호에 해당하는 경우 등 '중대한 공익상 필요가 있는 경우' — 일반 조항으로 범위가 매우 광범위하여 실효성에 대한 우려가 제기되고 있습니다. [S4]", 4)
    _numbered(doc, "**변호사-의뢰인 분쟁**: 변호사가 자신의 권리를 행사하거나 방어하기 위해 필요한 경우 (self-defense exception)", 5)
    _numbered(doc, "**다른 법률의 특별 규정**: 개별 법률에 특별한 규정이 있는 경우", 6)

    _heading3(doc, "미국")
    _numbered(doc, "**Voluntary Waiver**: 의뢰인의 의사에 의한 포기", 1)
    _numbered(doc, "**Implied/At-Issue Waiver**: 의뢰인이 소송에서 특권 대상 정보를 쟁점으로 삼은 경우 묵시적 포기", 2)
    _numbered(doc, "**Subject Matter Waiver**: 특정 커뮤니케이션의 자발적 공개 시 동일 주제에 관한 다른 커뮤니케이션에도 포기 확장", 3)
    _numbered(doc, "**Inadvertent Disclosure**: FRE 502(b) — 부주의 공개 시 합리적 조치를 취했다면 포기가 부정됩니다 (3요건: 부주의, 합리적 예방 조치, 신속한 시정 조치)", 4)
    _numbered(doc, "**Crime-Fraud Exception**: 의뢰인이 범죄·사기를 저지르거나 계획하면서 변호사의 서비스를 이용한 경우 — 변호사의 인식 여부와 무관합니다. [S5]", 5)
    _numbered(doc, "**Self-Defense Exception**: 변호사가 의뢰인의 주장에 대해 자기방어를 하는 경우", 6)
    _numbered(doc, "**Death of Client**: 의뢰인 사후에도 원칙적으로 존속하나, 유언 관련 분쟁 등 예외가 있습니다.", 7)

    _heading3(doc, "핵심 Gap")
    _bullet(doc, "한국법에는 **inadvertent disclosure에 대한 보호 규정**(FRE 502(b) 대응)이 없습니다. [Unverified]")
    _bullet(doc, "한국법의 '중대한 공익상 필요'는 미국의 crime-fraud exception보다 광범위하며, 수사기관의 확대 해석이 우려됩니다. [S4]")
    _bullet(doc, "한국법에는 **subject matter waiver**에 관한 규정이 없습니다. [Unverified]")
    _bullet(doc, "한국법에는 **at-issue waiver** (묵시적 포기)에 관한 규정이 없습니다. [Unverified]")

    # ── Axis 5 ──
    _heading2(doc, "5. 제3자 공유 시 특권 유지 여부")

    _heading3(doc, "한국")
    _body_paragraph(doc,
        "개정 제26조의2에는 제3자 공유 시 특권 유지에 관한 **명문 규정이 없습니다**. [Unverified] "
        "공동 이해관계인(common interest) 간 커뮤니케이션 공유 시 특권 유지 여부, 다수 의뢰인이 동일 변호사를 선임한 경우 "
        "(joint representation) 특권 처리 등은 향후 판례 또는 하위법령을 통해 해소되어야 할 것으로 판단됩니다."
    )

    _heading3(doc, "미국")
    _bullet(doc, "**Common Interest Doctrine**: 공동의 법적 이해관계를 가진 당사자들이 각자의 변호사를 통해 특권 정보를 공유하더라도, 공동 이해관계 합의(common interest agreement) 하에서 특권이 유지됩니다.")
    _bullet(doc, "**Joint Defense Privilege**: 형사사건에서 공동 피고인들이 방어 전략을 공유하는 경우 특권이 유지됩니다.")
    _bullet(doc, "**제3자 동석 원칙**: 불필요한 제3자에게 공개 시 원칙적으로 특권이 상실되나, 변호사 지시로 법률자문을 보조하는 제3자(통역사, 전문가 등)는 예외입니다. [S5, S6]")

    _heading3(doc, "핵심 Gap")
    _body_paragraph(doc,
        "한국법은 이 영역에서 **가장 큰 입법 공백**을 보이고 있습니다. 미국법의 Common Interest Doctrine에 해당하는 메커니즘이 "
        "전혀 없으며, 이는 복수 당사자 거래, 합작투자, 그룹 내 법률자문 공유 등에서 실무적 불확실성을 초래할 것으로 예상됩니다. [Unverified]"
    )

    # ── Axis 6 ──
    _heading2(doc, "6. 형사절차·수사기관 대응에서의 특권 행사 방법 및 한계")

    _heading3(doc, "한국")
    _body_paragraph(doc,
        "개정법의 핵심적 입법 동기는 수사기관의 **압수수색(押收搜索)** 과정에서 변호사-의뢰인 커뮤니케이션을 보호하는 데 있습니다. [A1] "
        "기존 실무에서는 변호사 사무실에 대한 압수수색 시 형사소송법 제112조(업무상 비밀과 압수)에 의한 제한적 보호만 "
        "존재하였습니다."
    )
    _bullet(doc, "**개정법 효과**: 법률자문 목적의 의사교환·서류에 대해 적극적 공개 거부권이 부여됩니다.")
    _bullet(doc, "**구체적 행사 절차** (예: 특권 주장의 방식, 선별 절차, 이의신청 절차 등)는 하위법령에 위임될 것으로 예상되나 현재 미확정입니다. [Unverified]")
    _bullet(doc, "수사기관이 '중대한 공익상 필요'를 근거로 압수를 강행할 가능성에 대한 우려가 존재합니다. [S4]")

    _heading3(doc, "미국")
    _bullet(doc, "**Privilege Log**: 특권 주장 측이 문서별 특권 주장 근거를 기재한 목록을 제출할 의무가 있습니다.")
    _bullet(doc, "**In Camera Review**: 법원이 비공개 심리를 통해 특권 해당 여부를 판단합니다.")
    _bullet(doc, "**Filter/Taint Team**: 수사기관 내 별도 팀이 압수 자료를 선별하여 특권 자료를 수사팀에서 격리합니다.")
    _bullet(doc, "**Federal Rule of Criminal Procedure 41(g)**: 부당 압수 재산의 반환 신청 절차가 마련되어 있습니다.")
    _bullet(doc, "**Motion to Quash / Protective Order**: 소환장·개시 명령에 대한 취소·보호 명령을 신청할 수 있습니다. [S5]")

    _heading3(doc, "핵심 Gap")
    _body_paragraph(doc,
        "한국법은 특권의 실체적 권리를 신설하였으나, 그 **절차적 행사 메커니즘** (privilege log, in camera review, taint team 등)이 "
        "부재합니다. 이는 하위법령 또는 대법원규칙을 통해 보충되어야 하며, 시행 전까지 가장 시급한 입법 과제로 판단됩니다. [Unverified]"
    )

    # ── Axis 7 ──
    _heading2(doc, "7. 위반 시 제재 및 구제수단")

    _heading3(doc, "한국")
    _body_paragraph(doc,
        "개정 제26조의2 자체에는 **위반 시 구체적 제재 규정이 포함되어 있지 않습니다**. [Unverified] "
        "관련 제재는 기존 법률에 분산되어 있습니다:"
    )
    _bullet(doc, "변호사법 제26조 위반 (비밀유지의무 위반): 징계사유에 해당합니다.")
    _bullet(doc, "형법 제317조 (업무상비밀누설죄): 3년 이하 징역 또는 700만원 이하 벌금에 처해집니다.")
    _bullet(doc, "특권 침해 시 증거능력 배제 여부: 형사소송법 제308조의2 위법수집증거배제법칙과의 관계에서 해석이 필요합니다. [Unverified]")

    _heading3(doc, "미국")
    _bullet(doc, "**증거배제(Exclusion/Suppression)**: 특권 위반으로 획득한 증거가 배제됩니다.")
    _bullet(doc, "**변호사 징계(Disciplinary Action)**: ABA Model Rule 1.6 위반 시 주 Bar Association에 의한 징계가 이루어집니다.")
    _bullet(doc, "**민사상 손해배상**: 비밀유지 위반에 대한 malpractice 소송이 가능합니다.")
    _bullet(doc, "**Disqualification**: 상대방 변호사가 특권 정보를 취득한 경우 해당 사건에서 자격이 박탈됩니다.")
    _bullet(doc, "**Contempt of Court**: 법원의 특권 인정 결정에도 불구하고 개시를 강행하는 경우 법원 모욕에 해당합니다. [S5]")

    _heading3(doc, "핵심 Gap")
    _body_paragraph(doc,
        "한국법은 특권 위반에 대한 독자적 제재 규정이 없어 기존 법률에 의존해야 하나, 이들은 '특권(privilege)' 침해가 아닌 "
        "'의무(duty)' 위반에 초점이 맞추어져 있어 보호에 한계가 있습니다. 특히 **위법수집증거배제**와 관련한 해석이 시급합니다. [Unverified]"
    )

    # ── Axis 8 ──
    _heading2(doc, "8. 시행 전 과도기 쟁점")

    _body_paragraph(doc, "*이 항목은 한국법에만 해당합니다 (미국법은 현행법으로 안정적으로 운영되고 있습니다).*")

    _heading3(doc, "(1) 시행일 이전 커뮤니케이션의 소급 보호")
    _body_paragraph(doc,
        "개정법 부칙은 **시행 이전의 의사교환 내용 및 서류·자료에도 적용**됨을 명시하고 있습니다. [S1, S2] "
        "이는 실무적으로 매우 중요한 의미를 가집니다:"
    )
    _bullet(doc, "시행 전에 이미 이루어진 압수수색으로 취득된 자료의 처리 문제가 발생합니다.")
    _bullet(doc, "시행 전 진행 중인 수사·재판에서의 증거 사용 가능 여부가 쟁점이 됩니다.")
    _bullet(doc, "기업은 시행 전부터 ACP 표시(privilege marking) 및 보호 실무를 확립해야 합니다. [S1]")

    _heading3(doc, "(2) 기존 압수수색 영장 실무와의 충돌")
    _body_paragraph(doc,
        "현행 형사소송법상 압수수색 영장 실무에서 변호사-의뢰인 커뮤니케이션에 대한 별도의 선별 절차가 체계화되어 있지 않습니다. "
        "개정법 시행 후 영장 집행 과정에서의 특권 주장 절차, 법원의 판단 기준 등이 구체화되어야 할 것입니다. [Unverified]"
    )

    _heading3(doc, "(3) 하위법령(대통령령·대법원규칙) 제정 동향")
    _body_paragraph(doc,
        "2026년 3월 현재, 하위법령 제정에 관한 구체적 일정이나 초안은 공개되지 않은 상태입니다. [Unverified] "
        "예상되는 하위법령 사항은 다음과 같습니다:"
    )
    _bullet(doc, "특권 주장의 절차적 방법 (privilege log에 해당하는 절차)")
    _bullet(doc, "수사기관의 선별(screening) 절차")
    _bullet(doc, "분쟁 시 법원의 판단 절차 (in camera review 등)")
    _bullet(doc, "사내변호사의 적용 범위 및 기준")

    # ── Axis 9 ──
    _heading2(doc, "9. 실무상 주요 쟁점 — 해석 공백 및 Gap 분석")

    _heading3(doc, "(1) 사내변호사(In-house Counsel) 문제")
    _body_paragraph(doc,
        "한국 개정법은 '변호사'라고만 규정하여 사내변호사의 커뮤니케이션이 특권 보호를 받는지 불명확합니다. 특히 법률자문과 "
        "경영판단이 혼합된 이메일·보고서의 처리가 문제됩니다. 미국에서는 'primary purpose test' 또는 'significant purpose test' "
        "(*In re Kellogg Brown & Root*)를 통해 dual-purpose communication의 특권 여부를 판단하고 있으나, 한국법에는 이러한 기준이 없습니다. [Unverified]"
    )

    _heading3(doc, "(2) 예외 조항의 광범성")
    _body_paragraph(doc,
        "'중대한 공익상 필요'라는 일반 조항이 수사기관에 의해 확대 해석될 우려가 있습니다. [S4] 미국의 crime-fraud exception은 "
        "\"(1) 의뢰인이 범죄·사기를 저지르거나 계획하고 있을 것, (2) 커뮤니케이션이 해당 범죄·사기를 진행하기 위한 것일 것\"이라는 "
        "비교적 구체적인 2요건 테스트를 적용하는 반면, 한국법의 예외 사유는 더 넓고 불확정적입니다."
    )

    _heading3(doc, "(3) 국제 소송에서의 상호 인정")
    _body_paragraph(doc,
        "한국에서 보호되는 자료가 미국 소송의 증거개시(discovery) 절차에서 어떻게 취급될지는 별도의 기준에 따릅니다. "
        "미국 법원은 일반적으로 comity 원칙과 choice-of-law 분석을 통해 외국의 privilege를 인정할지 판단하며, "
        "한국법의 성문화된 특권은 이 점에서 유리하게 작용할 수 있습니다. [S2]"
    )

    _heading3(doc, "(4) 전자적 커뮤니케이션의 보호 범위")
    _body_paragraph(doc,
        "개정법 제2항은 '전자적 형태'를 명시적으로 포함하였으나, 이메일·메신저·영상통화 등의 보호 범위, "
        "전자 포렌식(e-forensic) 이미징 시의 선별 절차, 클라우드 저장 자료에 대한 특권 적용 등은 구체화가 필요합니다. [Unverified]"
    )


def _add_glossary(doc):
    _heading1(doc, "IV. 용어 대조표 (Glossary)")

    headers = ["한국어", "영문 (English)", "비고"]
    rows = [
        ["비밀유지특권 / 비밀유지권", "Attorney-Client Privilege (ACP)", "개정 변호사법 제26조의2"],
        ["비밀유지의무", "Duty of Confidentiality", "기존 변호사법 제26조"],
        ["의사교환", "Communication", "제26조의2 제1항"],
        ["포기", "Waiver", ""],
        ["중대한 공익상 필요", "Significant Public Interest", "제26조의2 제3항 예외"],
        ["업무산출물 보호", "Work Product Doctrine", "제26조의2 제2항 유사 기능"],
        ["범죄·사기 예외", "Crime-Fraud Exception", ""],
        ["공동이해관계 법리", "Common Interest Doctrine", "한국법 규정 부재"],
        ["선별팀", "Filter / Taint Team", "한국법 규정 부재"],
        ["비밀문서 목록", "Privilege Log", "한국법 규정 부재"],
        ["비공개 심리", "In Camera Review", ""],
        ["사내변호사", "In-house Counsel", ""],
        ["위법수집증거배제", "Exclusionary Rule", "형사소송법 제308조의2"],
    ]
    _add_table(doc, headers, rows, col_widths_cm=[4.0, 5.5, 5.5])


def _add_bibliography(doc):
    _heading1(doc, "V. 주석부 참고문헌 (Annotated Bibliography)")

    _heading2(doc, "1차 자료 (Primary Sources)")
    _bullet(doc, "**[P1]** 변호사법 제26조의2 (신설), 국회 의결 2026.1.29., 국가법령정보센터 law.go.kr")
    _bullet(doc, "**[P2]** 변호사법 제26조 (현행), 국가법령정보센터 law.go.kr")
    _bullet(doc, "**[P3]** Federal Rules of Evidence, Rule 501, 28 U.S.C. App.")
    _bullet(doc, "**[P4]** Federal Rules of Evidence, Rule 502, 28 U.S.C. App.")
    _bullet(doc, "**[C1]** *Upjohn Co. v. United States*, 449 U.S. 383 (1981)")
    _bullet(doc, "**[A1]** 법무부 보도자료, \"국제 기준에 부합하는 비밀유지권 도입,\" 2026.1.29., korea.kr")

    _heading2(doc, "2차 자료 (Secondary Sources)")
    _bullet(doc, "**[S1]** 김·장 법률사무소, \"변호사와 의뢰인 간 비밀유지권(ACP)을 명문화한 변호사법 개정안 국회 본회의 통과,\" 2026 (kimchang.com)")
    _bullet(doc, "**[S2]** Baker McKenzie / Global Compliance News, \"South Korea: Introduction of Attorney-Client Privilege,\" 2026.3.2. (globalcompliancenews.com)")
    _bullet(doc, "**[S3]** 미주중앙일보, \"[한국법 이야기] 변호사-의뢰인 비밀유지권 한국 도입,\" 2026.3.3. (koreadaily.com)")
    _bullet(doc, "**[S4]** 이투데이, \"변호사 비밀유지권 생겼다지만…예외 조항 범위에 실효성은 '글쎄',\" 2026 (etoday.co.kr)")
    _bullet(doc, "**[S5]** DLA Piper, \"Legal Professional Privilege in United States,\" Global Privilege Guide (dlapiperintelligence.com)")
    _bullet(doc, "**[S6]** Baker McKenzie, \"Scope of Privilege — United States,\" Global Privilege Guide (resourcehub.bakermckenzie.com)")
    _bullet(doc, "**[S7]** 리걸타임즈, \"'변호사 비밀유지권 보장' 개정 변호사법 국회 통과\" (legaltimes.co.kr)")
    _bullet(doc, "**[S8]** 대한변호사협회 성명서, 2026.1.29. (naewoeilbo.com 재인용)")


def _add_verification_guide(doc):
    _heading1(doc, "VI. 검증 가이드 (Verification Guide)")

    headers = ["분석 내용", "검증 경로", "소스 ID"]
    rows = [
        ["제26조의2 조문 내용", "law.go.kr → 변호사법 → 개정이력", "P1"],
        ["기존 제26조 비밀유지의무", "law.go.kr → 변호사법 → 제26조", "P2"],
        ["FRE Rule 501 원문", "law.cornell.edu/rules/fre/rule_501", "P3"],
        ["FRE Rule 502 원문", "law.cornell.edu/rules/fre/rule_502", "P4"],
        ["Upjohn 판시 사항", "supreme.justia.com/cases/federal/us/449/383/", "C1"],
        ["소급 적용 여부", "개정법 부칙 확인 (law.go.kr)", "P1, S1"],
        ["예외 조항 광범성 우려", "etoday.co.kr 원문 기사", "S4"],
        ["사내변호사 적용 불명확", "조문상 명시 규정 부재 확인", "P1, S1"],
        ["하위법령 제정 동향", "법무부 보도자료, 국회 의안정보시스템", "A1"],
    ]
    _add_table(doc, headers, rows, col_widths_cm=[5.5, 6.5, 3.0])


def _add_unverified_summary(doc):
    _heading1(doc, "VII. 향후 확인 필요 사항 ([Unverified] 항목 종합)")

    items = [
        "사내변호사·외국법자문사에 대한 제26조의2 적용 여부 → 하위법령 또는 판례",
        "제3자 공유 시 특권 유지 메커니즘 (Common Interest Doctrine 등) → 하위법령 또는 판례",
        "특권 행사의 절차적 방법 (privilege log, screening, in camera review) → 대통령령·대법원규칙",
        "부주의 공개(inadvertent disclosure) 시 특권 보호 여부 → 판례",
        "Subject matter waiver / At-issue waiver 적용 여부 → 판례",
        "위반 시 독자적 제재 규정 → 하위법령",
        "위법수집증거배제법칙(형사소송법 제308조의2)과의 관계 → 판례",
        "시행 전 이미 압수된 자료의 소급 보호 구체적 범위 → 판례",
        "하위법령 제정 일정 및 초안 → 법무부·대법원 공개 예정",
    ]
    for i, item in enumerate(items, 1):
        _numbered(doc, item, i)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— 이상 —")
    _set_run_font(r, size=Pt(10), bold=True, color=COLOR_NAVY)


# ── BUILD ──────────────────────────────────────────────────────────────
def build():
    doc = Document()
    _setup_page(doc)
    _setup_header_footer(doc, "비밀유지특권 비교 분석  |  법무법인 진주  |  CONFIDENTIAL")
    _add_cover(doc)
    _add_toc(doc)
    _add_scope(doc)
    _add_comparative_matrix(doc)
    _add_detailed_analysis(doc)
    _add_glossary(doc)
    _add_bibliography(doc)
    _add_verification_guide(doc)
    _add_unverified_summary(doc)

    out_dir = Path(__file__).resolve().parents[1] / "output" / "reports"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "acp_comparison_kr_us_2026-03-09.docx"
    doc.save(str(out_path))
    print(f"Generated: {out_path}")
    return out_path


if __name__ == "__main__":
    build()
