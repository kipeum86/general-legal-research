import importlib.util
import json
import sys
from pathlib import Path

from docx import Document


MODULE_PATH = Path(__file__).resolve().parents[1] / "scripts" / "render_professional_legal_opinion_docx.py"
SPEC = importlib.util.spec_from_file_location("render_professional_legal_opinion_docx", MODULE_PATH)
renderer = importlib.util.module_from_spec(SPEC)
assert SPEC is not None and SPEC.loader is not None
sys.modules[SPEC.name] = renderer
SPEC.loader.exec_module(renderer)


def test_build_professional_docx_folds_explicit_audit_json(tmp_path: Path) -> None:
    body_claim = "Beta citation."
    md_text = "\n".join(
        [
            "# Formal Legal Opinion Letter",
            "- **Confidentiality**: Privileged",
            "- **Client**: Test Client",
            "- **Matter**: Test Matter",
            "- **As-of Date**: 2026-04-24",
            "- **Issue Date**: 2026-04-24",
            "- **Language**: Korean",
            "",
            "**Scope and Assumptions**",
            "",
            "Alpha.",
            "",
            body_claim,
        ]
    )
    md_path = tmp_path / "opinion.md"
    docx_path = tmp_path / "opinion.docx"
    audit_json = tmp_path / "audit.json"
    md_path.write_text(md_text, encoding="utf-8")
    start = md_text.index(body_claim)
    end = start + len(body_claim)
    audit_json.write_text(
        json.dumps(
            {
                "aggregated": [
                    {
                        "claim": {"text": body_claim, "sentence_span": {"start": start, "end": end}},
                        "verdict": {
                            "claim": {"text": body_claim, "sentence_span": {"start": start, "end": end}},
                            "label": "contradicted",
                            "verifier_name": "test-verifier",
                            "authority": 0.9,
                            "rationale": "Test contradiction.",
                            "evidence": [{"url": "https://example.test/source"}],
                        },
                    }
                ]
            }
        ),
        encoding="utf-8",
    )

    renderer.build_professional_docx(md_path, docx_path, audit_json=audit_json, output_dir=tmp_path)

    doc = Document(docx_path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Beta citation. [Unverified]" in text
    assert "부록: 검증 로그 (Citation Audit Log)" in text
    assert "Audit status: complete" in text
    assert str(audit_json) in text


def test_renderer_does_not_consume_latest_audit_without_explicit_request(tmp_path: Path) -> None:
    body_claim = "Beta citation."
    md_text = "\n".join(
        [
            "# Formal Legal Opinion Letter",
            "- **Client**: Test Client",
            "- **Matter**: Custom Matter",
            "",
            "**Scope and Assumptions**",
            "",
            body_claim,
        ]
    )
    md_path = tmp_path / "opinion.md"
    docx_path = tmp_path / "opinion.docx"
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    latest = output_dir / "citation-audit-latest.json"
    md_path.write_text(md_text, encoding="utf-8")
    start = md_text.index(body_claim)
    end = start + len(body_claim)
    latest.write_text(
        json.dumps(
            {
                "aggregated": [
                    {
                        "claim": {"text": body_claim, "sentence_span": {"start": start, "end": end}},
                        "verdict": {"label": "contradicted", "verifier_name": "test", "evidence": []},
                    }
                ]
            }
        ),
        encoding="utf-8",
    )

    renderer.build_professional_docx(md_path, docx_path, output_dir=output_dir)

    doc = Document(docx_path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "[Unverified]" not in text
    assert "Citation Audit Log" not in text


def test_parse_meta_uses_input_title_and_matter() -> None:
    lines = [
        "# Custom Opinion Title",
        "- **Matter**: Custom Matter",
        "- **Client**: Test Client",
        "",
        "**Scope and Assumptions**",
    ]

    meta, body_start = renderer.parse_meta(lines)

    assert meta.title == "Custom Opinion Title"
    assert meta.matter == "Custom Matter"
    assert body_start == 4
